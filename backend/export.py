import io
import os
import requests
from typing import List, Dict, Optional
from datetime import date, datetime

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image as XLImage

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                 Paragraph, Spacer, Image as RLImage, HRFlowable)
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Înregistrare fonturi cu suport complet Unicode/diacritice
def _try_register_fonts() -> tuple:
    """Try to register a Unicode-capable font; fallback to Helvetica."""
    # Candidate font sets: (regular, bold, name_prefix)
    candidates = [
        # macOS system fonts
        ("/System/Library/Fonts/Supplemental/Arial.ttf",
         "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
         "Arial"),
        # Linux: DejaVu (pre-installed on most distros including Railway/Debian)
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
         "DejaVu"),
        # Linux: Ubuntu fonts
        ("/usr/share/fonts/truetype/ubuntu/Ubuntu-R.ttf",
         "/usr/share/fonts/truetype/ubuntu/Ubuntu-B.ttf",
         "Ubuntu"),
        # Liberation fonts (RHEL/CentOS/Alpine)
        ("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
         "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
         "Liberation"),
    ]
    import os
    for reg, bold, prefix in candidates:
        if os.path.exists(reg) and os.path.exists(bold):
            try:
                pdfmetrics.registerFont(TTFont(prefix, reg))
                pdfmetrics.registerFont(TTFont(f"{prefix}-Bold", bold))
                return prefix, f"{prefix}-Bold"
            except Exception:
                pass
    return "Helvetica", "Helvetica-Bold"

_PDF_FONT, _PDF_FONT_BOLD = _try_register_fonts()

from PIL import Image as PILImage

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Risk thresholds ────────────────────────────────────────────────────
# VERY HIGH ≥ 91%  →  roșu
# HIGH      76–90% →  portocaliu
# MEDIUM    51–75% →  galben
# LOW       20–50% →  verde

def _risk_level(score: float) -> str:
    if score >= 90: return "very_high"
    if score >= 75: return "high"
    if score >= 60: return "medium"
    if score >= 45: return "low"
    return "low"

def _risk_label_ro(score: float) -> str:
    return {
        "very_high": "RISC FOARTE RIDICAT",
        "high":      "RISC RIDICAT",
        "medium":    "RISC MEDIU",
        "low":       "RISC SCĂZUT",
    }[_risk_level(score)]

# RGB foreground colors
_RISK_RGB = {
    "very_high": (192,  57,  43),   # roșu       #C0392B
    "high":      (175,  77,   0),   # portocaliu  #AF4D00
    "medium":    (154, 118,   0),   # galben      #9A7600
    "low":       ( 30, 132,  73),   # verde       #1E8449
}
# RGB background (tint)
_RISK_BG_RGB = {
    "very_high": (253, 236, 234),   # #FDECEA
    "high":      (254, 235, 210),   # #FEEBCF
    "medium":    (255, 249, 219),   # #FFF9DB
    "low":       (234, 250, 241),   # #EAFAF1
}

def _risk_hex_fg(score: float) -> str:
    r, g, b = _RISK_RGB[_risk_level(score)]
    return f"FF{r:02X}{g:02X}{b:02X}"

def _risk_color_pdf(score: float):
    r, g, b = _RISK_BG_RGB[_risk_level(score)]
    return colors.Color(r/255, g/255, b/255)


def _office_priority_map(offices: Optional[List[str]]) -> Dict[str, int]:
    cleaned: List[str] = []
    seen = set()
    for code in offices or []:
        c = (code or "").upper().strip()
        if not c or c in seen:
            continue
        seen.add(c)
        cleaned.append(c)
    preferred = [c for c in cleaned if c not in {"EM", "WO"}]
    primary = preferred[:1]
    secondary = preferred[1:]
    ordered = primary + ["EM", "WO"] + secondary
    return {code: idx for idx, code in enumerate(ordered)}


def _sortable_filing_date(entry: Dict):
    raw = (entry.get("applicationDate") or entry.get("registrationDate") or "").strip()
    if not raw:
        return date.max
    token = raw[:10]
    for fmt in ("%Y-%m-%d", "%d.%m.%Y", "%d/%m/%Y"):
        try:
            return datetime.strptime(token, fmt).date()
        except Exception:
            continue
    return date.max


def _sort_export_results(items: List[Dict], offices: Optional[List[str]]) -> List[Dict]:
    priority_map = _office_priority_map(offices)
    default_rank = len(priority_map) + 1

    return sorted(
        items or [],
        key=lambda entry: (
            priority_map.get((entry.get("office") or entry.get("tmOffice") or "").upper().strip(), default_rank),
            _sortable_filing_date(entry),
            -(entry.get("similarity", {}).get("combined_score", 0)),
        )
    )


def _risk_buckets(items: List[Dict]) -> Dict[str, List[Dict]]:
    buckets = {
        "very_high": [],
        "high": [],
        "medium": [],
        "low": [],
    }
    for item in items or []:
        level = _risk_level(item.get("similarity", {}).get("combined_score", 0))
        buckets[level].append(item)
    return buckets


_TMDN_BASE = "https://www.tmdn.org"
_IMG_HEADERS = {
    "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36",
    "Accept": "image/webp,image/apng,image/*,*/*;q=0.8",
    "Referer": "https://www.tmdn.org/tmview/",
    "Origin": "https://www.tmdn.org",
}


def _fetch_image_bytes(url: str, size=(60, 60)) -> Optional[bytes]:
    if not url:
        return None
    # URL-uri relative de la TMview → prefix cu domeniul
    if url.startswith("/"):
        url = _TMDN_BASE + url
    try:
        r = requests.get(url, timeout=10, headers=_IMG_HEADERS)
        if r.status_code == 200 and len(r.content) > 100:
            img = PILImage.open(io.BytesIO(r.content)).convert("RGBA")
            img.thumbnail(size, PILImage.LANCZOS)
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            buf.seek(0)
            return buf.read()
    except Exception:
        pass
    return None


def _fetch_local_logo_bytes() -> Optional[bytes]:
    candidates = [
        os.path.join(os.path.dirname(__file__), "..", "frontend", "protectmark-logo.png"),
        os.path.join(os.path.dirname(__file__), "..", "frontend", "protectmark-logo.svg"),
        os.path.join(os.path.dirname(__file__), "..", "protectmark-logo.png"),
        os.path.join(os.path.dirname(__file__), "..", "protectmark-logo.svg"),
    ]
    for path in candidates:
        if os.path.exists(path):
            try:
                with open(path, "rb") as f:
                    data = f.read()
                if data:
                    return data
            except Exception:
                pass
    return None


def _add_export_brand_header_pdf(story, query: str):
    logo_bytes = _fetch_local_logo_bytes()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "BrandHeaderTitle",
        parent=styles["Normal"],
        fontName=_PDF_FONT_BOLD,
        fontSize=16,
        leading=18,
        textColor=colors.HexColor("#0F3460"),
        alignment=TA_CENTER,
    )
    subtitle_style = ParagraphStyle(
        "BrandHeaderSubtitle",
        parent=styles["Normal"],
        fontName=_PDF_FONT_BOLD,
        fontSize=16,
        leading=18,
        textColor=colors.HexColor("#555555"),
        alignment=TA_CENTER,
    )
    logo_style = ParagraphStyle(
        "BrandHeaderLogo",
        parent=styles["Normal"],
        fontName=_PDF_FONT_BOLD,
        fontSize=14,
        leading=16,
        textColor=colors.HexColor("#0F3460"),
        alignment=TA_CENTER,
    )

    title = Paragraph("ProSearch", title_style)
    subtitle = Paragraph("RAPORT VERIFICARE DISPONIBILITATE MARCA", title_style)
    trademark = Paragraph(query, subtitle_style)

    if logo_bytes:
        try:
            logo = RLImage(io.BytesIO(logo_bytes), width=1.7*cm, height=1.7*cm)
        except Exception:
            logo = Paragraph("PM", logo_style)
    else:
        logo = Paragraph("PM", logo_style)

    header = Table([[logo, title], ["", subtitle], ["", trademark]], colWidths=[2.8*cm, 17.3*cm])
    header.setStyle(TableStyle([
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("SPAN", (0,0), (0,2)),
        ("ALIGN", (0,0), (-1,-1), "CENTER"),
        ("LEFTPADDING", (0,0), (-1,-1), 0),
        ("RIGHTPADDING", (0,0), (-1,-1), 0),
        ("TOPPADDING", (0,0), (-1,-1), 1),
        ("BOTTOMPADDING", (0,0), (-1,-1), 1),
        ("LEADING", (1,0), (1,2), 18),
    ]))
    story.append(header)
    story.append(Spacer(1, 0.35*cm))


def _add_export_brand_header_word(doc, query: str):
    logo_bytes = _fetch_local_logo_bytes()
    blue = RGBColor(0x0F, 0x34, 0x60)
    gray = RGBColor(0x55, 0x55, 0x55)
    tbl = doc.add_table(rows=3, cols=2)
    tbl.style = "Table Grid"
    _set_table_col_widths_cm(tbl, [2.8, 17.3])
    for cell in tbl.rows[0].cells:
        _set_cell_bg(cell, "FFFFFF")
    for row in tbl.rows:
        for cell in row.cells:
            _set_cell_bg(cell, "FFFFFF")

    c0 = tbl.cell(0,0).merge(tbl.cell(2,0))
    c1, c2, c3 = tbl.cell(0,1), tbl.cell(1,1), tbl.cell(2,1)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if logo_bytes:
        try:
            p0.add_run().add_picture(io.BytesIO(logo_bytes), width=Cm(2.3))
        except Exception:
            r = p0.add_run("PM")
            r.bold = True; r.font.size = Pt(14); r.font.name = "Arial"; r.font.color.rgb = blue
    else:
        r = p0.add_run("PM")
        r.bold = True; r.font.size = Pt(14); r.font.name = "Arial"; r.font.color.rgb = blue

    for cell in (c1, c2, c3):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("ProSearch")
    r1.bold = True; r1.font.size = Pt(15); r1.font.name = "Arial"; r1.font.color.rgb = blue

    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("RAPORT VERIFICARE DISPONIBILITATE MARCA")
    r2.bold = True; r2.font.size = Pt(15); r2.font.name = "Arial"; r2.font.color.rgb = blue

    p3 = c3.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(query)
    r3.bold = True; r3.font.size = Pt(15); r3.font.name = "Arial"; r3.font.color.rgb = blue

    _set_borders(tbl)
    for row in tbl.rows:
        row.height = Cm(0.95)
    doc.add_paragraph()


# ── Excel ──────────────────────────────────────────────────────────────
def build_excel(query: str, nice_classes: List[str], offices: List[str],
                results: List[Dict], similar: List[Dict] = None,
                expired_conflicts: List[Dict] = None, expired_similar: List[Dict] = None) -> bytes:
    from datetime import datetime as _dt

    def _xdate(d):
        if not d: return ""
        try: return _dt.strptime(str(d)[:10], "%Y-%m-%d").strftime("%d.%m.%Y")
        except: return str(d)[:10]

    wb = Workbook()
    ws = wb.active
    ws.title = "Raport Similaritate"

    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left   = Alignment(horizontal="left",   vertical="center", wrap_text=True)
    mid    = Alignment(horizontal="left",   vertical="center", wrap_text=True)
    thin   = Side(style="thin", color="FFD0D0D0")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # ── Titlu ──────────────────────────────────────────────────────────
    NCOLS = 13
    ws.merge_cells(f"A1:{get_column_letter(NCOLS)}1")
    c = ws["A1"]
    c.value     = f"Raport Cercetare Disponibilitate Marca: {query}"
    c.font      = Font(bold=True, size=13, color="FF0F3460")
    c.alignment = center
    ws.row_dimensions[1].height = 26

    active_results = _sort_export_results((results or []) + (similar or []), offices)
    expired_results = _sort_export_results((expired_conflicts or []) + (expired_similar or []), offices)
    total_results = len(active_results) + len(expired_results)

    ws.merge_cells(f"A2:{get_column_letter(NCOLS)}2")
    c = ws["A2"]
    c.value = (f"Clase NICE: {', '.join(nice_classes)}  |  Teritorii: {', '.join(offices)}  |  "
               f"Rezultate active: {len(active_results)}  |  Mărci expirate: {len(expired_results)}  |  "
               f"Total rezultate: {total_results}  |  Data: {date.today().strftime('%d.%m.%Y')}")
    c.font      = Font(italic=True, size=9, color="FF555555")
    c.alignment = center
    ws.row_dimensions[2].height = 16
    ws.row_dimensions[3].height = 4

    # ── Antet coloane ─────────────────────────────────────────────────
    # Col: 1=#  2=Nivel risc  3=Scor  4=Sigla  5=Denumire marca  6=Birou
    #      7=Status  8=Titular  9=Data depunere  10=Data inregistrare
    #      11=Data expirare  12=Clase NICE  13=Produse/Servicii
    headers = [
        "#", "Nivel risc", "Scor", "Sigla",
        "Denumire marca", "Birou / Oficiu", "Status",
        "Titular / Solicitant",
        "Data depunere", "Data inregistrare", "Data expirare",
        "Clase NICE", "Produse si servicii",
    ]
    col_widths = [4, 16, 7, 9, 20, 18, 12, 24, 12, 13, 12, 10, 28]

    hdr_font = Font(bold=True, color="FFFFFFFF", size=9)
    hdr_fill = PatternFill("solid", fgColor="FF0F3460")
    for col, (h, w) in enumerate(zip(headers, col_widths), 1):
        cell = ws.cell(row=4, column=col, value=h)
        cell.font      = hdr_font
        cell.fill      = hdr_fill
        cell.alignment = center
        cell.border    = border
        ws.column_dimensions[get_column_letter(col)].width = w
    ws.row_dimensions[4].height = 28

    # Configure worksheet so the full table fits on an A4 landscape page
    # and remains readable without manual resizing in Excel.
    ws.freeze_panes = "A5"
    ws.sheet_view.zoomScale = 85
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.page_margins.left = 0.25
    ws.page_margins.right = 0.25
    ws.page_margins.top = 0.5
    ws.page_margins.bottom = 0.5
    ws.page_margins.header = 0.2
    ws.page_margins.footer = 0.2
    ws.print_options.horizontalCentered = True
    ws.print_title_rows = "1:4"

    # ── Rânduri date ──────────────────────────────────────────────────
    RISK_LABELS_XL = {
        "very_high": "RISC FOARTE RIDICAT",
        "high":      "RISC RIDICAT",
        "medium":    "RISC MEDIU",
        "low":       "RISC SCAZUT",
    }

    def add_result_rows(items: List[Dict], row_start: int, display_start: int, expired_section: bool = False):
        for offset, tm in enumerate(items):
            sim    = tm.get("similarity", {})
            score  = sim.get("combined_score", 0)
            lvl    = _risk_level(score)
            row_idx = row_start + offset
            row_no = display_start + offset

            applicant = ", ".join(a.get("name", "") for a in tm.get("applicants", []) if a.get("name")) or "—"
            nice_nums = ", ".join(f"Cls {c}" for c in tm.get("niceClass") or [])
            if tm.get("goodAndServices"):
                nice_desc = "\n".join(
                    f"Cls {g['niceClass']}: {g['goodsAndServices']}"
                    for g in tm["goodAndServices"] if g.get("goodsAndServices")
                )
            else:
                nice_desc = "; ".join(nd.get("short", "") for nd in tm.get("niceDetailed") or [])

            status = tm.get("status", "") or "—"
            exp_date = _xdate(tm.get("expiryDate", ""))
            exp_note = f"{exp_date} *" if exp_date and not tm.get("expiryIsReal") else (exp_date or "—")

            row_vals = [
                row_no,
                "MARCĂ EXPIRATĂ" if expired_section else RISK_LABELS_XL[lvl],
                f"{score}%",
                "",
                tm.get("tmName", "—"),
                tm.get("officeName", tm.get("office", "—")),
                status,
                applicant,
                _xdate(tm.get("applicationDate", "")),
                _xdate(tm.get("registrationDate", "")),
                exp_note,
                nice_nums,
                nice_desc,
            ]

            r, g, b = _RISK_RGB[lvl]
            rb, gb, bb = _RISK_BG_RGB[lvl]
            fg_hex = f"FF{r:02X}{g:02X}{b:02X}"
            bg_hex = "FFFDECEA" if expired_section else f"FF{rb:02X}{gb:02X}{bb:02X}"
            risk_fill  = PatternFill("solid", fgColor=bg_hex)
            white_fill = PatternFill("solid", fgColor="FFFFFFFF")
            alt_fill   = PatternFill("solid", fgColor="FFFDF7F7") if expired_section else (PatternFill("solid", fgColor="FFF7F9FC") if row_no % 2 == 0 else white_fill)

            for col, val in enumerate(row_vals, 1):
                cell = ws.cell(row=row_idx, column=col, value=val)
                cell.border = border

                if col == 1:
                    cell.alignment = center
                    cell.fill      = white_fill
                    cell.font      = Font(size=9, color="FF888888")
                elif col == 2:
                    cell.alignment = center
                    cell.fill      = risk_fill
                    cell.font      = Font(bold=True, size=9, color="FFC0392B" if expired_section else fg_hex)
                elif col == 3:
                    cell.alignment = center
                    cell.fill      = risk_fill
                    cell.font      = Font(bold=True, size=12, color="FFC0392B" if expired_section else fg_hex)
                elif col == 4:
                    cell.alignment = center
                    cell.fill      = alt_fill
                    cell.font      = Font(size=14, color="FFBBBBBB")
                    cell.value     = "TM"
                elif col == 5:
                    cell.alignment = left
                    cell.fill      = alt_fill
                    cell.font      = Font(bold=True, size=10, color="FF1a1a2e")
                elif col in (6, 7):
                    cell.alignment = center
                    cell.fill      = alt_fill
                    cell.font      = Font(size=9)
                elif col == 8:
                    cell.alignment = left
                    cell.fill      = alt_fill
                    cell.font      = Font(size=9, bold=True)
                elif col in (9, 10):
                    cell.alignment = center
                    cell.fill      = alt_fill
                    cell.font      = Font(size=9)
                elif col == 11:
                    cell.alignment = center
                    cell.fill      = alt_fill
                    cell.font      = Font(size=9, bold=True, color="FFC0392B" if exp_date else "FF888888")
                elif col == 12:
                    cell.alignment = center
                    cell.fill      = alt_fill
                    cell.font      = Font(size=8, color="FF0F3460")
                else:
                    cell.alignment = left
                    cell.fill      = alt_fill
                    cell.font      = Font(size=8, color="FF444444")

            max_text_len = max(len(str(row_vals[4] or "")), len(str(row_vals[7] or "")), len(str(row_vals[12] or "")))
            estimated_lines = max(2, min(8, (max_text_len // 38) + 1))
            ws.row_dimensions[row_idx].height = max(52, estimated_lines * 16)

            img_bytes = _fetch_image_bytes(tm.get("imageUrl"), size=(48, 48))
            if img_bytes:
                try:
                    xl_img = XLImage(io.BytesIO(img_bytes))
                    xl_img.width = 42; xl_img.height = 42
                    ws.add_image(xl_img, f"D{row_idx}")
                    ws.cell(row=row_idx, column=4).value = ""
                except Exception:
                    pass

    active_header_row = 5
    ws.merge_cells(f"A{active_header_row}:M{active_header_row}")
    active_cell = ws.cell(row=active_header_row, column=1, value=f"Mărci active ({len(active_results)})")
    active_cell.font = Font(bold=True, size=14, color="FF0F3460")
    active_cell.fill = PatternFill("solid", fgColor="FFE8F0FB")
    active_cell.alignment = center
    for col in range(1, NCOLS + 1):
        ws.cell(row=active_header_row, column=col).border = border

    add_result_rows(active_results, active_header_row + 1, 1)

    next_row = active_header_row + 1 + len(active_results)
    if expired_results:
        ws.merge_cells(f"A{next_row}:M{next_row}")
        section_cell = ws.cell(row=next_row, column=1, value=f"Mărci expirate ({len(expired_results)})")
        section_cell.font = Font(bold=True, size=14, color="FFC0392B")
        section_cell.fill = PatternFill("solid", fgColor="FFFDECEA")
        section_cell.alignment = center
        for col in range(1, NCOLS + 1):
            ws.cell(row=next_row, column=col).border = border
        add_result_rows(expired_results, next_row + 1, 1, expired_section=True)
        next_row = next_row + 1 + len(expired_results)

    # ── Nota subsol ────────────────────────────────────────────────────
    fn_row = next_row + 1
    ws.merge_cells(f"A{fn_row}:{get_column_letter(NCOLS)}{fn_row}")
    fn = ws.cell(row=fn_row, column=1,
                 value="* Data expirare marcata cu * este estimata (inregistrare + 10 ani). Datele confirmate de TMview nu au asterisc.")
    fn.font      = Font(italic=True, size=8, color="FF888888")
    fn.alignment = left

    # ── Foaie concluzii ───────────────────────────────────────────────
    risk_buckets = _risk_buckets(active_results + expired_results)
    very_high = risk_buckets["very_high"]
    high = risk_buckets["high"]
    medium = risk_buckets["medium"]
    low = risk_buckets["low"]

    high_risk_count = len(very_high) + len(high)
    medium_risk_count = len(medium)
    minimal_risk_count = len(low)

    if high_risk_count:
        risk_heading = "Concluzie privind riscul relativ: risc semnificativ"
        risk_fill_hex = "FDECEA"
        risk_font_hex = "C0392B"
        risk_body = (
            f"Au fost identificate {high_risk_count} mărci cu risc ridicat sau foarte ridicat, "
            f"{medium_risk_count} mărci cu risc mediu și {minimal_risk_count} mărci cu risc minim/scăzut. "
            "În raport cu motivele relative de refuz, acest rezultat indică o probabilitate relevantă de opoziție "
            "sau de refuz întemeiat pe existența unor drepturi anterioare, în special atunci când se cumulează "
            "similitudinea semnului cu proximitatea produselor sau serviciilor revendicate."
        )
    elif medium_risk_count:
        risk_heading = "Concluzie privind riscul relativ: risc moderat"
        risk_fill_hex = "FEEBCF"
        risk_font_hex = "AF4D00"
        risk_body = (
            f"Au fost identificate {high_risk_count} mărci cu risc ridicat sau foarte ridicat, "
            f"{medium_risk_count} mărci cu risc mediu și {minimal_risk_count} mărci cu risc minim/scăzut. "
            "În raport cu motivele relative de refuz, rezultatele indică necesitatea unei evaluări juridice "
            "suplimentare a mărcilor aflate în zona medie, deoarece riscul de confuzie depinde de impresia de "
            "ansamblu a semnelor și de gradul de apropiere dintre produsele sau serviciile vizate."
        )
    else:
        risk_heading = "Concluzie privind riscul relativ: risc redus"
        risk_fill_hex = "EAFAF1"
        risk_font_hex = "1E8449"
        risk_body = (
            f"Au fost identificate {high_risk_count} mărci cu risc ridicat sau foarte ridicat, "
            f"{medium_risk_count} mărci cu risc mediu și {minimal_risk_count} mărci cu risc minim/scăzut. "
            "Pe baza conflictelor identificate, motivele relative de refuz par limitate; totuși, disponibilitatea "
            "juridică finală trebuie confirmată prin analiza individuală a drepturilor anterioare relevante."
        )

    conclusion_blocks = [
        (risk_heading, risk_body, risk_fill_hex, risk_font_hex),
        (
            "Opinia privind motivele absolute de refuz",
            "Prezentul raport are în principal o funcție de cercetare a conflictelor relative și nu poate stabili în mod definitiv incidența motivelor absolute de refuz. În practica oficiilor și a jurisprudenței europene, examinarea motivelor absolute vizează în special caracterul distinctiv, eventuala descriptivitate, caracterul uzual, posibilul caracter înșelător și conformitatea semnului cu ordinea publică. În consecință, marca trebuie analizată separat și prin raportare exactă la lista produselor și serviciilor revendicate.",
            "E8F0FB",
            "0F3460",
        ),
        (
            "Opinia privind validitatea mărcii",
            "Condiția validității unei viitoare înregistrări depinde cumulativ de absența impedimentelor absolute, de inexistența unor drepturi anterioare opozabile și de formularea adecvată a specificației de produse și servicii. Chiar și în ipoteza unui risc relativ redus, validitatea nu poate fi considerată automat îndeplinită fără verificarea completă a cadrului juridic aplicabil pe teritoriile selectate.",
            "E8F0FB",
            "0F3460",
        ),
        (
            "Opinia privind distinctivitatea",
            "Distinctivitatea trebuie apreciată din perspectiva publicului relevant și în raport direct cu produsele sau serviciile pentru care se solicită protecția. În mod obișnuit, semnele fanteziste sau arbitrare au o forță distinctivă mai ridicată, în timp ce semnele descriptive, laudative ori slab individualizante sunt mai expuse obiecțiilor la examinare și beneficiază de o protecție mai restrânsă. Dacă denumirea propusă evocă în mod imediat natura, calitatea, destinația sau alte caracteristici ale produselor sau serviciilor, se recomandă o reevaluare a semnului înainte de depunere.",
            "E8F0FB",
            "0F3460",
        ),
        (
            "Recomandare finală",
            "Rezultatele acestui raport trebuie utilizate ca instrument de triere a riscului, nu ca opinie juridică definitivă. Pentru depunere, este recomandată o analiză specializată privind opozabilitatea drepturilor anterioare, formularea claselor NICE și sustenabilitatea mărcii sub aspectul distinctivității și al motivelor absolute de refuz.",
            "E8F0FB",
            "0F3460",
        ),
    ]

    ws2 = wb.create_sheet("Concluzii")
    ws2.column_dimensions["A"].width = 28
    ws2.column_dimensions["B"].width = 110
    ws2.merge_cells("A1:B1")
    ws2["A1"] = f"Concluzii raport disponibilitate marcă: {query}"
    ws2["A1"].font = Font(bold=True, size=14, color="FF0F3460")
    ws2["A1"].alignment = left
    ws2.merge_cells("A2:B2")
    ws2["A2"] = (
        f"Clase NICE: {', '.join(nice_classes)}  |  Teritorii: {', '.join(offices)}  |  "
        f"Rezultate active: {len(active_results)}  |  Mărci expirate: {len(expired_results)}  |  "
        f"Total rezultate: {total_results}  |  Data: {date.today().strftime('%d.%m.%Y')}"
    )
    ws2["A2"].font = Font(italic=True, size=9, color="FF555555")
    ws2["A2"].alignment = left

    ws2["A4"] = "Sinteză risc"
    ws2["A4"].font = Font(bold=True, size=11, color="FF0F3460")
    ws2["A5"] = "Mărci cu risc ridicat / foarte ridicat"
    ws2["B5"] = high_risk_count
    ws2["A6"] = "Mărci cu risc mediu"
    ws2["B6"] = medium_risk_count
    ws2["A7"] = "Mărci cu risc minim / scăzut"
    ws2["B7"] = minimal_risk_count
    ws2["A8"] = "Mărci expirate"
    ws2["B8"] = len(expired_results)
    for row_idx in (5, 6, 7, 8):
        ws2[f"A{row_idx}"].font = Font(bold=True, size=10, color="FF333333")
        ws2[f"B{row_idx}"].font = Font(bold=True, size=10, color="FF0F3460")
        ws2[f"A{row_idx}"].fill = PatternFill("solid", fgColor="FFF7F9FC")
        ws2[f"B{row_idx}"].fill = PatternFill("solid", fgColor="FFF7F9FC")
        ws2[f"A{row_idx}"].border = border
        ws2[f"B{row_idx}"].border = border

    row_ptr = 10
    for title, body, fill_hex, font_hex in conclusion_blocks:
        ws2.merge_cells(f"A{row_ptr}:B{row_ptr}")
        ws2[f"A{row_ptr}"] = title
        ws2[f"A{row_ptr}"].font = Font(bold=True, size=11, color=f"FF{font_hex}")
        ws2[f"A{row_ptr}"].fill = PatternFill("solid", fgColor=f"FF{fill_hex}")
        ws2[f"A{row_ptr}"].border = border
        ws2[f"A{row_ptr}"].alignment = left
        row_ptr += 1
        ws2.merge_cells(f"A{row_ptr}:B{row_ptr + 2}")
        ws2[f"A{row_ptr}"] = body
        ws2[f"A{row_ptr}"].font = Font(size=10, color="FF444444")
        ws2[f"A{row_ptr}"].alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        for inner in range(row_ptr, row_ptr + 3):
            ws2[f"A{inner}"].border = border
        row_ptr += 4

    ws2.freeze_panes = "A5"
    ws2.page_setup.orientation = ws2.ORIENTATION_LANDSCAPE
    ws2.page_setup.paperSize = ws2.PAPERSIZE_A4
    ws2.page_setup.fitToWidth = 1
    ws2.page_setup.fitToHeight = 0
    ws2.page_margins.left = 0.25
    ws2.page_margins.right = 0.25
    ws2.page_margins.top = 0.5
    ws2.page_margins.bottom = 0.5
    ws2.page_margins.header = 0.2
    ws2.page_margins.footer = 0.2
    ws2.print_options.horizontalCentered = True

    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return buf.read()



# ── PDF ────────────────────────────────────────────────────────────────
def build_pdf(query: str, nice_classes: List[str], offices: List[str],
              results: List[Dict], similar: List[Dict] = None,
              expired_conflicts: List[Dict] = None, expired_similar: List[Dict] = None) -> bytes:
    from reportlab.lib.pagesizes import landscape, A4
    from reportlab.platypus import KeepTogether, PageBreak
    from datetime import datetime as dt

    PAGE = landscape(A4)
    LM = RM = 1.4 * cm
    TM = BM = 1.4 * cm
    W  = PAGE[0] - LM - RM

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=PAGE,
                            leftMargin=LM, rightMargin=RM,
                            topMargin=TM, bottomMargin=BM)
    styles = getSampleStyleSheet()
    story  = []

    _add_export_brand_header_pdf(story, query)

    BLUE   = colors.HexColor("#0F3460")
    DKGRAY = colors.HexColor("#444444")
    LGRAY  = colors.HexColor("#F7F9FC")

    def sty(name, **kw):
        return ParagraphStyle(name, parent=styles["Normal"], fontName=_PDF_FONT, **kw)
    def styb(name, **kw):
        return ParagraphStyle(name, parent=styles["Normal"], fontName=_PDF_FONT_BOLD, **kw)
    def ctbl(*args, **kwargs):
        tbl = Table(*args, **kwargs)
        tbl.hAlign = "CENTER"
        return tbl

    RISK_PDF_COLORS = {
        "very_high": (colors.HexColor("#FDECEA"), colors.HexColor("#C0392B")),
        "high":      (colors.HexColor("#FEEBCF"), colors.HexColor("#AF4D00")),
        "medium":    (colors.HexColor("#FFF9DB"), colors.HexColor("#9A7600")),
        "low":       (colors.HexColor("#EAFAF1"), colors.HexColor("#1E8449")),
    }
    RISK_LABELS_RO = {
        "very_high": "RISC FOARTE RIDICAT",
        "high":      "RISC RIDICAT",
        "medium":    "RISC MEDIU",
        "low":       "RISC SCĂZUT",
    }

    active_results = _sort_export_results((results or []) + (similar or []), offices)
    expired_results = _sort_export_results((expired_conflicts or []) + (expired_similar or []), offices)
    all_results = active_results + expired_results
    risk_buckets = _risk_buckets(all_results)
    very_high = risk_buckets["very_high"]
    high      = risk_buckets["high"]
    medium    = risk_buckets["medium"]
    low       = risk_buckets["low"]

    def fmt_date(d):
        if not d: return "—"
        try: return dt.strptime(d[:10], "%Y-%m-%d").strftime("%d.%m.%Y")
        except: return d[:10] if d else "—"

    risky_count  = len(very_high) + len(high)
    similar_count = len(medium) + len(low)
    safe = risky_count == 0

    # Distributie pe oficii din rezultate
    geo_counts = {}
    for tm in all_results:
        o = tm.get("office") or tm.get("tmOffice") or "?"
        geo_counts[o] = geo_counts.get(o, 0) + 1
    geo_sorted = sorted(geo_counts.items(), key=lambda x: x[1], reverse=True)
    geo_max = geo_sorted[0][1] if geo_sorted else 1

    # ─── COVER PAGE (dashboard ca in UI) ─────────────────────────────
    story.append(Paragraph(
        "Verificare Disponibilitate Marcă",
        sty("app_lbl", fontSize=10, textColor=DKGRAY, spaceAfter=4)
    ))
    story.append(Paragraph(
        query,
        styb("q_title", fontSize=22, textColor=colors.HexColor("#1a1a2e"), spaceAfter=6)
    ))
    story.append(HRFlowable(width="100%", thickness=2, color=BLUE))
    story.append(Spacer(1, 0.5*cm))

    # Badges — 3 coloane x 2 randuri (ca in UI)
    CW3 = [W/3 - 0.2*cm] * 3
    GAP = TableStyle([
        ("VALIGN",        (0,0), (-1,-1), "MIDDLE"),
        ("LEFTPADDING",   (0,0), (-1,-1), 4),
        ("RIGHTPADDING",  (0,0), (-1,-1), 4),
        ("TOPPADDING",    (0,0), (-1,-1), 3),
        ("BOTTOMPADDING", (0,0), (-1,-1), 3),
    ])

    def badge_cell(text, val, bg_hex, fg_hex, bold_val=True):
        bg_c = colors.HexColor(bg_hex)
        fg_c = colors.HexColor(fg_hex)
        lbl  = Paragraph(text, sty(f"bl{text[:6]}", fontSize=7.5, textColor=fg_c))
        num  = Paragraph(str(val), styb(f"bv{text[:6]}", fontSize=14, textColor=fg_c, leading=16) if bold_val
                         else sty(f"bv{text[:6]}", fontSize=8.5, textColor=fg_c, leading=12))
        cell_tbl = ctbl([[lbl], [num]], colWidths=[W/3 - 0.6*cm])
        cell_tbl.setStyle(TableStyle([
            ("BACKGROUND",    (0,0), (-1,-1), bg_c),
            ("TOPPADDING",    (0,0), (-1,-1), 8),
            ("BOTTOMPADDING", (0,0), (-1,-1), 8),
            ("LEFTPADDING",   (0,0), (-1,-1), 10),
            ("RIGHTPADDING",  (0,0), (-1,-1), 10),
            ("BOX",           (0,0), (-1,-1), 1, fg_c),
        ]))
        return cell_tbl

    row1 = [
        badge_cell("Total găsite",                  len(all_results),  "#E8F0FB", "#0F3460"),
        badge_cell("Risc ridicat / f.ridicat >=70%", risky_count,
                   "#FDECEA" if not safe else "#EAFAF1",
                   "#C0392B" if not safe else "#1E8449"),
        badge_cell("Risc mediu / scăzut 40-70%",   similar_count,     "#FFF3CD", "#856404"),
    ]
    row2 = [
        badge_cell("Clase NICE",    ", ".join(nice_classes),            "#E8F0FB", "#0F3460", bold_val=False),
        badge_cell("Mărci expirate", str(len(expired_results)),            "#FDECEA", "#C0392B"),
        badge_cell("Data raport",   date.today().strftime("%d.%m.%Y"),  "#F2F3F4", "#566573", bold_val=False),
    ]

    for row in [row1, row2]:
        bt = ctbl([row], colWidths=[W/3, W/3, W/3])
        bt.setStyle(GAP)
        story.append(bt)
        story.append(Spacer(1, 0.25*cm))

    # Distribuție pe oficii (ca secțiunea Geo din UI)
    if geo_sorted:
        story.append(Spacer(1, 0.3*cm))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#D0D7E3")))
        story.append(Spacer(1, 0.3*cm))
        story.append(Paragraph(
            "Distribuție pe oficii",
            styb("geo_title", fontSize=9, textColor=BLUE, spaceAfter=6)
        ))

        BAR_W = 4 * cm
        geo_hdr = [
            Paragraph("Cod", styb("gh0", fontSize=7, textColor=colors.white)),
            Paragraph("Oficiu", styb("gh1", fontSize=7, textColor=colors.white)),
            Paragraph("Mărci", styb("gh2", fontSize=7, textColor=colors.white)),
            Paragraph("Distribuție", styb("gh3", fontSize=7, textColor=colors.white)),
        ]
        geo_rows = [geo_hdr]
        geo_style = [
            ("BACKGROUND",    (0,0), (-1,0), BLUE),
            ("FONTSIZE",      (0,0), (-1,-1), 7),
            ("ALIGN",         (0,0), (-1,-1), "CENTER"),
            ("ALIGN",         (1,1), (1,-1), "LEFT"),
            ("VALIGN",        (0,0), (-1,-1), "MIDDLE"),
            ("GRID",          (0,0), (-1,-1), 0.3, colors.HexColor("#D0D7E3")),
            ("TOPPADDING",    (0,0), (-1,-1), 3),
            ("BOTTOMPADDING", (0,0), (-1,-1), 3),
            ("LEFTPADDING",   (0,0), (-1,-1), 5),
        ]
        OFFICE_NAMES_SHORT = {
            "EM":"EUIPO (UE)", "WO":"WIPO (Intl)", "RO":"OSIM Romania",
            "DE":"DPMA Germania", "FR":"INPI Franta", "IT":"UIBM Italia",
            "ES":"OEPM Spania",  "PL":"UPRP Polonia","BG":"BPO Bulgaria",
            "HU":"HIPO Ungaria", "CZ":"IPO Cehia",   "AT":"APO Austria",
            "NL":"BOIP Olanda",  "BE":"BOIP Belgia",  "PT":"INPI Portugalia",
            "SE":"PRV Suedia",   "DK":"DKPTO Danemarca","GB":"UKIPO Marea Britanie",
        }
        for ri, (code, cnt) in enumerate(geo_sorted, 1):
            pct  = cnt / geo_max
            bar_fill = BAR_W * pct
            is_max = cnt == geo_max
            bar_tbl = ctbl(
                [[""]], colWidths=[bar_fill]
            )
            bar_tbl.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,-1),
                 colors.HexColor("#C0392B") if is_max else BLUE),
                ("TOPPADDING",    (0,0), (-1,-1), 4),
                ("BOTTOMPADDING", (0,0), (-1,-1), 4),
            ]))
            bar_wrap = ctbl([[bar_tbl, ""]], colWidths=[bar_fill, BAR_W - bar_fill])
            bar_wrap.setStyle(TableStyle([
                ("BACKGROUND", (1,0), (1,0), colors.HexColor("#E9ECEF")),
                ("TOPPADDING",    (0,0), (-1,-1), 0),
                ("BOTTOMPADDING", (0,0), (-1,-1), 0),
                ("LEFTPADDING",   (0,0), (-1,-1), 0),
                ("RIGHTPADDING",  (0,0), (-1,-1), 0),
            ]))
            name_short = OFFICE_NAMES_SHORT.get(code, code)
            geo_rows.append([
                Paragraph(code, styb(f"gc{ri}", fontSize=7,
                          textColor=colors.HexColor("#C0392B") if is_max else BLUE)),
                Paragraph(name_short, sty(f"gn{ri}", fontSize=7, textColor=DKGRAY)),
                Paragraph(str(cnt), styb(f"gct{ri}", fontSize=8,
                          textColor=colors.HexColor("#C0392B") if is_max else BLUE)),
                bar_wrap,
            ])
            if ri % 2 == 0:
                geo_style.append(("BACKGROUND", (0,ri), (-1,ri), LGRAY))

        CGEO = [1.2*cm, W - 1.2*cm - 2*cm - BAR_W, 2*cm, BAR_W]
        geo_tbl = ctbl(geo_rows, colWidths=CGEO, repeatRows=1)
        geo_tbl.setStyle(TableStyle(geo_style))
        story.append(geo_tbl)

    story.append(PageBreak())

    # ─── STRATEGY PAGE ────────────────────────────────────────────────

    # ─── RESULTS SECTION ──────────────────────────────────────────────
    story.append(Paragraph(
        f"Rezultate analiza similaritate — {query}",
        styb("rh", fontSize=13, textColor=BLUE, spaceAfter=4)
    ))
    story.append(Paragraph(
        f"Active: {len(active_results)} mărci  |  Expirate: {len(expired_results)}  |  Clase NICE: {', '.join(nice_classes)}  |  Data: {date.today().strftime('%d.%m.%Y')}",
        sty("rsub", fontSize=8, textColor=DKGRAY, spaceAfter=12)
    ))

    if not active_results and not expired_results:
        story.append(Paragraph("Niciun conflict detectat.", styb("nc0", fontSize=11, textColor=colors.HexColor("#1E8449"))))
        doc.build(story)
        buf.seek(0)
        return buf.read()

    story.append(Paragraph(
        f"Mărci active ({len(active_results)})",
        styb("acth", fontSize=15, textColor=BLUE, spaceAfter=8)
    ))

    # Column widths  (landscape A4 cu margini 1.4cm → W ≈ 812pt)
    STRIP = 0.35 * cm   # strip colorat stânga
    LOGO  = 2.60 * cm   # logo marcă — mai mare
    SCORE = 4.20 * cm   # coloana scor — mai lată
    INFO  = W - STRIP - LOGO - SCORE

    MAX_GS = 600  # caractere max / clasă G&S

    for i, tm in enumerate(active_results):
        sim   = tm.get("similarity") or {}
        score = sim.get("combined_score") or 0
        lvl   = _risk_level(score)
        bg_c, fg_c = RISK_PDF_COLORS[lvl]
        risk_label = RISK_LABELS_RO[lvl]

        # ── Logo ────────────────────────────────────────────────────────
        img_bytes = _fetch_image_bytes(tm.get("imageUrl"), size=(70, 70))
        if img_bytes:
            try:
                logo_el = RLImage(io.BytesIO(img_bytes), width=1.90*cm, height=1.90*cm)
            except Exception:
                logo_el = Paragraph("™", sty(f"lf{i}", fontSize=28, alignment=TA_CENTER, textColor=colors.HexColor("#CCCCCC")))
        else:
            logo_el = Paragraph("™", sty(f"le{i}", fontSize=28, alignment=TA_CENTER, textColor=colors.HexColor("#CCCCCC")))

        # ── Status ──────────────────────────────────────────────────────
        status = tm.get("status") or ""
        sl     = status.lower()
        if "registered" in sl:
            stat_txt, stat_fg = "✔ Înregistrată", "#1E8449"
        elif "filed" in sl or "pending" in sl:
            stat_txt, stat_fg = "⏳ Depusă", "#B7950B"
        elif any(w in sl for w in ("expir","lapsed","cancelled","refused","withdrawn")):
            stat_txt, stat_fg = "✖ Expirată/Anulată", "#C0392B"
        else:
            stat_txt, stat_fg = status or "—", "#666666"

        office      = tm.get("office") or ""
        office_name = tm.get("officeName") or ""
        applicant   = ", ".join(a.get("name","") for a in (tm.get("applicants") or []) if a.get("name")) or "—"
        app_addr    = "; ".join(a.get("address","") for a in (tm.get("applicants") or []) if a.get("address"))
        reps        = ", ".join((r.get("fullName") or r.get("organizationName",""))
                                for r in (tm.get("representatives") or [])
                                if r.get("fullName") or r.get("organizationName"))
        an          = tm.get("applicationNumber") or "—"
        rn          = tm.get("registrationNumber") or "—"
        exp_str     = fmt_date(tm.get("expiryDate") or "")
        exp_mark    = " *" if exp_str and not tm.get("expiryIsReal") else ""
        is_multi    = sim.get("is_multiword", False)

        # ── NICE chips (sortate crescator) ──────────────────────────────
        nice_detailed = sorted(
            tm.get("niceDetailed") or [],
            key=lambda nd: int(str(nd.get("class", 0))) if str(nd.get("class","0")).isdigit() else 0
        )
        if nice_detailed:
            nice_html = "  ·  ".join(
                f'<font color="#0F3460"><b>Cls {nd["class"]}</b></font>'
                f'<font color="#555555"> – {nd.get("short") or ""}</font>'
                for nd in nice_detailed
            )
        else:
            nice_html = "  ".join(
                f'<font color="#0F3460"><b>Cls {c}</b></font>'
                for c in sorted(tm.get("niceClass") or [],
                                key=lambda x: int(x) if str(x).isdigit() else 0)
            )

        # ── Info column ──────────────────────────────────────────────────
        name_p = Paragraph(
            tm.get("tmName") or "—",
            styb(f"nm{i}", fontSize=16, textColor=colors.HexColor("#1a1a2e"), leading=19, spaceAfter=7)
        )
        meta_p = Paragraph(
            f'<font color="#0F3460" size="9"><b> {office} </b></font>'
            f'<font color="#888888" size="8">  {office_name}  </font>'
            f'<font color="{stat_fg}" size="8.5"><b>{stat_txt}</b></font>'
            + (f'  <font color="#6C3483" size="8"><b>[multi-cuvânt]</b></font>' if is_multi else ""),
            sty(f"mt{i}", leading=13, spaceAfter=7)
        )
        owner_p = Paragraph(
            f'<font color="#999999" size="7.5">Titular</font><br/>'
            f'<font color="#222222" size="9.5"><b>{applicant}</b></font>',
            sty(f"ow{i}", leading=13, spaceAfter=6)
        )
        nums_p = Paragraph(
            f'<font color="#999999" size="7">Nr. marcă: </font><font color="#1a1a2e" size="8">{rn}</font>'
            f'<font color="#CCCCCC">   |   </font>'
            f'<font color="#999999" size="7">Nr. depozit: </font><font color="#1a1a2e" size="8">{an}</font>',
            sty(f"nr{i}", leading=11, spaceAfter=5)
        )
        d_app = fmt_date(tm.get("applicationDate") or "")
        d_reg = fmt_date(tm.get("registrationDate") or "")
        dates_p = Paragraph(
            f'<font color="#999999" size="7">Depus: </font><font color="#1a1a2e" size="8">{d_app or "—"}</font>'
            f'<font color="#CCCCCC">   |   </font>'
            f'<font color="#999999" size="7">Înreg.: </font><font color="#1a1a2e" size="8">{d_reg or "—"}</font>'
            f'<font color="#CCCCCC">   |   </font>'
            f'<font color="#999999" size="7">Expiră: </font>'
            f'<font color="#C0392B" size="8"><b>{exp_str or "—"}{exp_mark}</b></font>',
            sty(f"dt{i}", leading=11, spaceAfter=7)
        )
        nice_p = Paragraph(
            nice_html or "—",
            sty(f"nc{i}", fontSize=8, leading=12, spaceAfter=0,
                backColor=colors.HexColor("#EEF3FB"))
        )

        info_cell = [name_p, meta_p, owner_p, nums_p, dates_p, nice_p]

        # ── Score column ─────────────────────────────────────────────────
        t_score = sim.get("textual_score") or 0
        p_score = sim.get("phonetic_score") or 0
        bar_w   = SCORE - 0.6*cm
        t_frac  = min(t_score * 0.70, 100) / 100
        p_frac  = min(p_score * 0.30, max(0, 100 - t_score * 0.70)) / 100
        e_frac  = max(0.0, 1.0 - t_frac - p_frac)
        t_bw    = max(bar_w * t_frac, 1)
        p_bw    = max(bar_w * p_frac, 0.5)
        e_bw    = max(bar_w * e_frac, 0.5)

        seg_bar = ctbl([["", "", ""]], colWidths=[t_bw, p_bw, e_bw])
        seg_bar.setStyle(TableStyle([
            ("BACKGROUND",    (0,0),(0,0), colors.HexColor("#2980B9")),
            ("BACKGROUND",    (1,0),(1,0), colors.HexColor("#8E44AD")),
            ("BACKGROUND",    (2,0),(2,0), colors.HexColor("#D0D7E3")),
            ("TOPPADDING",    (0,0),(-1,-1), 5),
            ("BOTTOMPADDING", (0,0),(-1,-1), 5),
            ("LEFTPADDING",   (0,0),(-1,-1), 0),
            ("RIGHTPADDING",  (0,0),(-1,-1), 0),
        ]))

        inner_w = SCORE - 0.5*cm
        risk_badge_tbl = ctbl(
            [[Paragraph(risk_label, styb(f"rb{i}", fontSize=7.5, textColor=fg_c,
                                         alignment=TA_CENTER, leading=10))]],
            colWidths=[inner_w]
        )
        risk_badge_tbl.setStyle(TableStyle([
            ("BACKGROUND",    (0,0),(-1,-1), bg_c),
            ("BOX",           (0,0),(-1,-1), 0.8, fg_c),
            ("TOPPADDING",    (0,0),(-1,-1), 5),
            ("BOTTOMPADDING", (0,0),(-1,-1), 5),
            ("LEFTPADDING",   (0,0),(-1,-1), 4),
            ("RIGHTPADDING",  (0,0),(-1,-1), 4),
        ]))

        half_w = inner_w / 2
        seg_labels = ctbl([[
            Paragraph(f"📝 {t_score}%", sty(f"tsl{i}", fontSize=7, textColor=colors.HexColor("#2980B9"), alignment=TA_CENTER)),
            Paragraph(f"🔊 {p_score}%", sty(f"psl{i}", fontSize=7, textColor=colors.HexColor("#8E44AD"), alignment=TA_CENTER)),
        ]], colWidths=[half_w, half_w])
        seg_labels.setStyle(TableStyle([
            ("LEFTPADDING",   (0,0),(-1,-1), 0),
            ("RIGHTPADDING",  (0,0),(-1,-1), 0),
            ("TOPPADDING",    (0,0),(-1,-1), 2),
            ("BOTTOMPADDING", (0,0),(-1,-1), 2),
        ]))

        score_cell = [
            Paragraph(f"{score}%", styb(f"sc{i}", fontSize=32, textColor=fg_c,
                                        alignment=TA_CENTER, leading=34, spaceAfter=6)),
            risk_badge_tbl,
            Spacer(1, 0.28*cm),
            seg_bar,
            seg_labels,
            Spacer(1, 0.18*cm),
            Paragraph(f"Jaro-Winkler: {sim.get('jaro_winkler',0)}%",
                      sty(f"jw{i}", fontSize=6.5, textColor=DKGRAY, alignment=TA_CENTER)),
            Paragraph(f"Levenshtein: {sim.get('levenshtein_distance',0)} car.",
                      sty(f"lv{i}", fontSize=6.5, textColor=DKGRAY, alignment=TA_CENTER)),
        ]

        # ── Asamblare card ───────────────────────────────────────────────
        card_data = [["", logo_el, info_cell, score_cell]]
        card_tbl  = ctbl(card_data, colWidths=[STRIP, LOGO, INFO, SCORE])
        card_tbl.setStyle(TableStyle([
            ("BACKGROUND",    (0,0), (0,0), fg_c),                    # strip colorat
            ("BACKGROUND",    (1,0), (2,0), colors.HexColor("#FAFCFF")), # corp foarte deschis
            ("BACKGROUND",    (3,0), (3,0), bg_c),                    # scor: tint risc
            ("VALIGN",        (0,0), (-1,-1), "MIDDLE"),
            ("VALIGN",        (2,0), (2,0),   "TOP"),
            # strip: fără padding
            ("TOPPADDING",    (0,0), (0,0), 0),
            ("BOTTOMPADDING", (0,0), (0,0), 0),
            ("LEFTPADDING",   (0,0), (0,0), 0),
            ("RIGHTPADDING",  (0,0), (0,0), 0),
            # logo
            ("TOPPADDING",    (1,0), (1,0), 16),
            ("BOTTOMPADDING", (1,0), (1,0), 16),
            ("LEFTPADDING",   (1,0), (1,0), 12),
            ("RIGHTPADDING",  (1,0), (1,0), 10),
            # info
            ("TOPPADDING",    (2,0), (2,0), 16),
            ("BOTTOMPADDING", (2,0), (2,0), 16),
            ("LEFTPADDING",   (2,0), (2,0), 10),
            ("RIGHTPADDING",  (2,0), (2,0), 16),
            # scor
            ("TOPPADDING",    (3,0), (3,0), 16),
            ("BOTTOMPADDING", (3,0), (3,0), 16),
            ("LEFTPADDING",   (3,0), (3,0), 10),
            ("RIGHTPADDING",  (3,0), (3,0), 10),
            # border card
            ("BOX",           (0,0), (-1,-1), 1.5, colors.HexColor("#C8D5EA")),
            ("LINEBEFORE",    (3,0), (3,0),   0.8, colors.HexColor("#D8E3F0")),
        ]))

        # ── Detalii suplimentare ─────────────────────────────────────────
        def det(label, value):
            if not value:
                return None
            return Paragraph(
                f'<font size="6" color="#999999">{label}</font><br/>'
                f'<font size="7.5" color="#1a1a2e"><b>{str(value)}</b></font>',
                sty(f"d{i}{label[:3]}", leading=11, spaceAfter=2)
            )

        pub_date  = tm.get("publicationDate") or ""
        opp_start = tm.get("oppositionStartDate") or ""
        opp_end   = tm.get("oppositionEndDate") or ""
        mark_feat = " · ".join(filter(None,[tm.get("markFeature") or "", tm.get("kindMark") or ""]))
        vienna    = ", ".join(tm.get("viennaCodes") or [])
        desig     = ", ".join(tm.get("designatedCountries") or [])
        found_by  = tm.get("_found_by") or ""
        exp_note  = "* Data estimata (inreg. + 10 ani)" if tm.get("expiryDate") and not tm.get("expiryIsReal") else ""

        extra_fields = [det(lbl, val) for lbl, val in [
            ("Data publicare (450)",    fmt_date(pub_date)),
            ("Perioada opozitie",       f"{fmt_date(opp_start)} – {fmt_date(opp_end)}" if opp_start else ""),
            ("Natura marcii (550)",     mark_feat),
            ("Coduri Vienna (531)",     vienna),
            ("Tari desemnate Madrid",   desig),
            ("Reprezentant (740)",      reps),
            ("Adresa titular",          app_addr),
            ("ST13",                    tm.get("ST13") or ""),
            ("Gasit prin varianta",     found_by),
            ("Nota expirare",           exp_note),
        ] if val]
        extra_fields = [f for f in extra_fields if f is not None]

        detail_elements = []
        if extra_fields:
            ncols = 3
            rows  = []
            for j in range(0, len(extra_fields), ncols):
                chunk = list(extra_fields[j:j+ncols])
                while len(chunk) < ncols:
                    chunk.append("")
                rows.append(chunk)
            det_tbl = ctbl(rows, colWidths=[W/ncols]*ncols)
            det_tbl.setStyle(TableStyle([
                ("BACKGROUND",    (0,0), (-1,-1), colors.HexColor("#FAFBFD")),
                ("TOPPADDING",    (0,0), (-1,-1), 5),
                ("BOTTOMPADDING", (0,0), (-1,-1), 5),
                ("LEFTPADDING",   (0,0), (-1,-1), 10),
                ("BOX",           (0,0), (-1,-1), 0.5, colors.HexColor("#E0E7F0")),
                ("INNERGRID",     (0,0), (-1,-1), 0.3, colors.HexColor("#E8EDF5")),
            ]))
            detail_elements.append(det_tbl)

        # ── G&S: colectam toate clasele, sortam crescator, facem blocuri ──
        # Construim un dict {nc_str: {text, short, desc}} din ambele surse
        all_cls: dict = {}
        for g in (tm.get("goodAndServices") or []):
            nc = str(g.get("niceClass") or "")
            if nc and nc.isdigit():
                if nc not in all_cls:
                    all_cls[nc] = {"text": "", "short": "", "desc": ""}
                all_cls[nc]["text"]  = g.get("goodsAndServices") or ""
                all_cls[nc]["short"] = g.get("niceShort") or ""
        for nd in nice_detailed:
            nc = str(nd.get("class", ""))
            if nc and nc.isdigit():
                if nc not in all_cls:
                    all_cls[nc] = {"text": "", "short": "", "desc": ""}
                if not all_cls[nc]["short"]:
                    all_cls[nc]["short"] = nd.get("short") or ""
                all_cls[nc]["desc"] = nd.get("description") or ""

        gs_blocks = []
        for nc in sorted(all_cls.keys(), key=lambda x: int(x)):
            info  = all_cls[nc]
            short = info["short"]
            text  = info["text"]
            desc  = info["desc"]

            # Titlu: doar numărul clasei (short apare deja în chips-urile din card)
            box_rows = [
                [Paragraph(f"Clasa {nc}", styb(f"gt{i}{nc}", fontSize=8.5, textColor=BLUE,
                                               leading=11, spaceAfter=0))],
            ]
            # Short ca primă linie de conținut (o singură apariție)
            if short:
                box_rows.append([Paragraph(
                    short, sty(f"gts{i}{nc}", fontSize=8, textColor=DKGRAY,
                               leading=11, spaceAfter=0)
                )])
            if text:
                disp = text[:MAX_GS] + ("…" if len(text) > MAX_GS else "")
                box_rows.append([Paragraph(
                    disp, sty(f"gtx{i}{nc}", fontSize=8, textColor=colors.HexColor("#444444"),
                              leading=12, spaceAfter=0)
                )])
            if desc:
                box_rows.append([Paragraph(
                    desc, sty(f"gdc{i}{nc}", fontSize=7, leading=11, spaceAfter=0,
                              textColor=colors.HexColor("#AAAAAA"))
                )])

            box = ctbl(box_rows, colWidths=[W])
            box.setStyle(TableStyle([
                ("BACKGROUND",    (0,0), (-1,-1), colors.white),
                ("TOPPADDING",    (0,0), (0,0),   7),
                ("BOTTOMPADDING", (0,-1),(0,-1),  8),
                ("TOPPADDING",    (0,1), (-1,-1),  4),
                ("BOTTOMPADDING", (0,0), (-1,-2),  4),
                ("LEFTPADDING",   (0,0), (-1,-1), 14),
                ("RIGHTPADDING",  (0,0), (-1,-1), 12),
                ("LINEBEFORE",    (0,0), (0,-1),   3, BLUE),
                ("BOX",           (0,0), (-1,-1), 0.5, colors.HexColor("#D8E3F0")),
            ]))
            gs_blocks.append(box)
            gs_blocks.append(Spacer(1, 0.28*cm))

        story.append(KeepTogether([card_tbl] + detail_elements))
        if gs_blocks:
            story.append(Paragraph(
                '<font color="#0F3460"><b>CLASIFICARE INTERNATIONALA NISA / PRODUSE SI SERVICII</b></font>',
                sty(f"gsh{i}", fontSize=9, spaceAfter=5, spaceBefore=7)
            ))
            for el in gs_blocks:
                story.append(el)
        story.append(Spacer(1, 0.70*cm))

    if expired_results:
        exp_bold7_s = styb("expB7s", fontSize=7, leading=9)
        exp_cell7_s = sty("expC7s", fontSize=7, leading=9)
        story.append(PageBreak())
        story.append(Paragraph(
            f"Mărci expirate ({len(expired_results)})",
            styb("exph", fontSize=15, textColor=colors.HexColor("#C0392B"), spaceAfter=6)
        ))
        story.append(Paragraph(
            "Aceste rezultate sunt afișate separat deoarece au statut expirat/inactiv.",
            sty("expsub", fontSize=8, textColor=DKGRAY, spaceAfter=10)
        ))

        exp_cols = ["#", "Denumire marcă", "Oficiu", "Titular", "Status", "Scor"]
        exp_w = [0.6*cm, 5*cm, 3*cm, 8*cm, 5*cm, 2*cm]
        exp_data = [[Paragraph(h, exp_bold7_s) for h in exp_cols]]
        for ri, tm in enumerate(expired_results, 1):
            score2 = tm.get("similarity",{}).get("combined_score",0)
            applicant = ", ".join(a.get("name","") for a in tm.get("applicants",[]) if a.get("name")) or "—"
            exp_data.append([
                Paragraph(str(ri), exp_cell7_s),
                Paragraph(tm.get("tmName") or "—", exp_bold7_s),
                Paragraph(tm.get("office") or "—", exp_cell7_s),
                Paragraph(applicant, exp_cell7_s),
                Paragraph(tm.get("status") or "—", exp_cell7_s),
                Paragraph(f"{score2}%", styb(f"expsc{ri}", fontSize=7, textColor=colors.HexColor("#C0392B"))),
            ])
        exp_tbl = ctbl(exp_data, colWidths=exp_w, repeatRows=1)
        exp_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#C0392B")),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("FONTNAME", (0,0), (-1,0), _PDF_FONT_BOLD),
            ("FONTSIZE", (0,0), (-1,-1), 7),
            ("ALIGN", (0,0), (-1,-1), "CENTER"),
            ("ALIGN", (1,1), (1,-1), "LEFT"),
            ("ALIGN", (3,1), (3,-1), "LEFT"),
            ("ALIGN", (4,1), (4,-1), "LEFT"),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.HexColor("#FFF5F5"), colors.white]),
            ("GRID", (0,0), (-1,-1), 0.3, colors.HexColor("#E6B8B8")),
            ("TOPPADDING", (0,0), (-1,-1), 2),
            ("BOTTOMPADDING", (0,0), (-1,-1), 2),
            ("LEFTPADDING", (0,0), (-1,-1), 3),
        ]))
        story.append(exp_tbl)
        story.append(Spacer(1, 0.6*cm))

    # ─── SUMMARY PAGE ──────────────────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph("Sumar rezultate", styb("SH", fontSize=16, textColor=BLUE, spaceAfter=10)))

    bold7_s = styb("B7s", fontSize=7, leading=9)
    cell7_s = sty("C7s", fontSize=7, leading=9)

    sum_groups = [
        (very_high, "Risc foarte ridicat  (>= 90%)", "#C0392B", "#FDECEA"),
        (high,      "Risc ridicat  (75-89%)",         "#AF4D00", "#FEEBCF"),
        (medium,    "Risc mediu  (60-74%)",            "#9A7600", "#FFF9DB"),
        (low,       "Risc scăzut  (45-59%)",           "#1E8449", "#EAFAF1"),
    ]

    for grp, lbl, fg_hex, bg_hex in sum_groups:
        fgc = colors.HexColor(fg_hex)
        bgc = colors.HexColor(bg_hex)
        story.append(ctbl(
            [[Paragraph(f"  {lbl}  -  {len(grp)} mărci", styb(f"sh{lbl[:4]}", fontSize=10, textColor=fgc))]],
            colWidths=[W],
            style=[("BACKGROUND",(0,0),(0,0),bgc),("TOPPADDING",(0,0),(0,0),6),
                   ("BOTTOMPADDING",(0,0),(0,0),6),("LEFTPADDING",(0,0),(0,0),8),
                   ("GRID",(0,0),(-1,-1),0.3,colors.HexColor("#D0D7E3"))]))
        story.append(Spacer(1, 0.2*cm))

        if not grp:
            story.append(Paragraph("  Nicio marcă.", sty("nt", fontSize=8, textColor=DKGRAY)))
            story.append(Spacer(1, 0.3*cm))
            continue

        sum_cols = ["#", "Denumire marcă", "Oficiu", "Titular", "Status", "Scor"]
        sum_w    = [0.6*cm, 5*cm, 3*cm, 7*cm, 3*cm, 2*cm]
        sum_data = [[Paragraph(h, bold7_s) for h in sum_cols]]
        for ri, tm in enumerate(grp, 1):
            score2    = tm.get("similarity",{}).get("combined_score",0)
            applicant = ", ".join(a.get("name","") for a in tm.get("applicants",[]) if a.get("name")) or "n/a"
            _, fg2    = RISK_PDF_COLORS[_risk_level(score2)]
            sum_data.append([
                Paragraph(str(ri),              cell7_s),
                Paragraph(tm.get("tmName") or "n/a", bold7_s),
                Paragraph(tm.get("office") or "n/a", cell7_s),
                Paragraph(applicant,              cell7_s),
                Paragraph(tm.get("status") or "n/a", cell7_s),
                Paragraph(f"{score2}%", styb(f"sc2{ri}", fontSize=7, textColor=fg2)),
            ])
        sum_tbl = ctbl(sum_data, colWidths=sum_w, repeatRows=1)
        sum_tbl.setStyle(TableStyle([
            ("BACKGROUND",    (0,0), (-1,0), BLUE),
            ("TEXTCOLOR",     (0,0), (-1,0), colors.white),
            ("FONTNAME",      (0,0), (-1,0), _PDF_FONT_BOLD),
            ("FONTSIZE",      (0,0), (-1,-1), 7),
            ("ALIGN",         (0,0), (-1,-1), "CENTER"),
            ("ALIGN",         (1,1), (1,-1), "LEFT"),
            ("ALIGN",         (3,1), (3,-1), "LEFT"),
            ("ALIGN",         (4,1), (4,-1), "LEFT"),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white, LGRAY]),
            ("GRID",          (0,0), (-1,-1), 0.3, colors.HexColor("#D0D7E3")),
            ("TOPPADDING",    (0,0), (-1,-1), 2),
            ("BOTTOMPADDING", (0,0), (-1,-1), 2),
            ("LEFTPADDING",   (0,0), (-1,-1), 3),
        ]))
        story.append(sum_tbl)
        story.append(Spacer(1, 0.4*cm))

    story.append(ctbl(
        [[Paragraph(f"  Mărci expirate  -  {len(expired_results)} mărci", styb("shexp", fontSize=10, textColor=colors.HexColor("#C0392B")))]],
        colWidths=[W],
        style=[("BACKGROUND",(0,0),(0,0),colors.HexColor("#FDECEA")),("TOPPADDING",(0,0),(0,0),6),
               ("BOTTOMPADDING",(0,0),(0,0),6),("LEFTPADDING",(0,0),(0,0),8),
               ("GRID",(0,0),(-1,-1),0.3,colors.HexColor("#D0D7E3"))]))
    story.append(Spacer(1, 0.2*cm))
    if expired_results:
        exp_sum_data = [[Paragraph(h, bold7_s) for h in ["#", "Denumire marcă", "Oficiu", "Status", "Scor"]]]
        for ri, tm in enumerate(expired_results, 1):
            score2 = tm.get("similarity",{}).get("combined_score",0)
            exp_sum_data.append([
                Paragraph(str(ri), cell7_s),
                Paragraph(tm.get("tmName") or "—", bold7_s),
                Paragraph(tm.get("office") or "—", cell7_s),
                Paragraph(tm.get("status") or "—", cell7_s),
                Paragraph(f"{score2}%", styb(f"exsum{ri}", fontSize=7, textColor=colors.HexColor("#C0392B"))),
            ])
        exp_sum_tbl = ctbl(exp_sum_data, colWidths=[0.6*cm, 9.5*cm, 3*cm, 8.2*cm, 2*cm], repeatRows=1)
        exp_sum_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#C0392B")),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.HexColor("#FFF5F5"), colors.white]),
            ("GRID", (0,0), (-1,-1), 0.3, colors.HexColor("#E6B8B8")),
            ("ALIGN", (0,0), (-1,-1), "CENTER"),
            ("ALIGN", (1,1), (1,-1), "LEFT"),
            ("ALIGN", (3,1), (3,-1), "LEFT"),
        ]))
        story.append(exp_sum_tbl)
    else:
        story.append(Paragraph("  Nicio marcă expirată.", sty("ntexp", fontSize=8, textColor=DKGRAY)))
    story.append(Spacer(1, 0.4*cm))

    # ─── CONCLUSIONS PAGE ──────────────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph("Concluzii și recomandări", styb("CH", fontSize=16, textColor=BLUE, spaceAfter=10)))

    high_risk_count = len(very_high) + len(high)
    medium_risk_count = len(medium)
    minimal_risk_count = len(low)

    if high_risk_count:
        risk_heading = "Concluzie privind riscul relativ: risc semnificativ"
        risk_color = colors.HexColor("#C0392B")
        risk_body = (
            f"Au fost identificate {high_risk_count} mărci cu risc ridicat sau foarte ridicat, "
            f"{medium_risk_count} mărci cu risc mediu și {minimal_risk_count} mărci cu risc minim/scăzut. "
            "În raport cu motivele relative de refuz, acest rezultat indică o probabilitate relevantă de "
            "opoziție sau de refuz întemeiat pe existența unor drepturi anterioare, în special atunci când "
            "se cumulează similitudinea semnului cu proximitatea produselor sau serviciilor revendicate."
        )
    elif medium_risk_count:
        risk_heading = "Concluzie privind riscul relativ: risc moderat"
        risk_color = colors.HexColor("#AF4D00")
        risk_body = (
            f"Au fost identificate {high_risk_count} mărci cu risc ridicat sau foarte ridicat, "
            f"{medium_risk_count} mărci cu risc mediu și {minimal_risk_count} mărci cu risc minim/scăzut. "
            "În raport cu motivele relative de refuz, rezultatele indică necesitatea unei evaluări juridice "
            "suplimentare a mărcilor aflate în zona medie, deoarece riscul de confuzie depinde de impresia de "
            "ansamblu a semnelor și de gradul de apropiere dintre produsele sau serviciile vizate."
        )
    else:
        risk_heading = "Concluzie privind riscul relativ: risc redus"
        risk_color = colors.HexColor("#1E8449")
        risk_body = (
            f"Au fost identificate {high_risk_count} mărci cu risc ridicat sau foarte ridicat, "
            f"{medium_risk_count} mărci cu risc mediu și {minimal_risk_count} mărci cu risc minim/scăzut. "
            "Pe baza conflictelor identificate, motivele relative de refuz par limitate; totuși, disponibilitatea "
            "juridică finală trebuie confirmată prin analiza individuală a drepturilor anterioare relevante."
        )

    conclusion_blocks = [
        (risk_heading, risk_body, risk_color),
        (
            "Opinia privind motivele absolute de refuz",
            "Prezentul raport are în principal o funcție de cercetare a conflictelor relative și nu poate stabili "
            "în mod definitiv incidența motivelor absolute de refuz. În practica oficiilor și a jurisprudenței "
            "europene, examinarea motivelor absolute vizează în special caracterul distinctiv, eventuala "
            "descriptivitate, caracterul uzual, posibilul caracter înșelător și conformitatea semnului cu ordinea "
            "publică. În consecință, marca trebuie analizată separat și prin raportare exactă la lista produselor "
            "și serviciilor revendicate.",
            BLUE,
        ),
        (
            "Opinia privind validitatea mărcii",
            "Condiția validității unei viitoare înregistrări depinde cumulativ de absența impedimentelor absolute, "
            "de inexistența unor drepturi anterioare opozabile și de formularea adecvată a specificației de produse "
            "și servicii. Chiar și în ipoteza unui risc relativ redus, validitatea nu poate fi considerată automat "
            "îndeplinită fără verificarea completă a cadrului juridic aplicabil pe teritoriile selectate.",
            BLUE,
        ),
        (
            "Opinia privind distinctivitatea",
            "Distinctivitatea trebuie apreciată din perspectiva publicului relevant și în raport direct cu produsele "
            "sau serviciile pentru care se solicită protecția. În mod obișnuit, semnele fanteziste sau arbitrare au "
            "o forță distinctivă mai ridicată, în timp ce semnele descriptive, laudative ori slab individualizante "
            "sunt mai expuse obiecțiilor la examinare și beneficiază de o protecție mai restrânsă. Dacă denumirea "
            "propusă evocă în mod imediat natura, calitatea, destinația sau alte caracteristici ale produselor sau "
            "serviciilor, se recomandă o reevaluare a semnului înainte de depunere.",
            BLUE,
        ),
        (
            "Recomandare finală",
            "Rezultatele acestui raport trebuie utilizate ca instrument de triere a riscului, nu ca opinie juridică "
            "definitivă. Pentru depunere, este recomandată o analiză specializată privind opozabilitatea drepturilor "
            "anterioare, formularea claselor NICE și sustenabilitatea mărcii sub aspectul distinctivității și al "
            "motivelor absolute de refuz.",
            BLUE,
        ),
    ]

    for idx, (title, body, title_color) in enumerate(conclusion_blocks, 1):
        story.append(Paragraph(title, styb(f"ct{idx}", fontSize=10.5, textColor=title_color, spaceAfter=4)))
        story.append(Paragraph(body, sty(f"cb{idx}", fontSize=8.8, leading=12.5, textColor=DKGRAY, spaceAfter=8)))

    doc.build(story)
    buf.seek(0)
    return buf.read()


# ── Word helpers ──────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _rgb(score: float) -> RGBColor:
    r, g, b = _RISK_RGB[_risk_level(score)]
    return RGBColor(r, g, b)


def _bg_hex(score: float) -> str:
    r, g, b = _RISK_BG_RGB[_risk_level(score)]
    return f"{r:02X}{g:02X}{b:02X}"



def _fmt_date(d: str) -> str:
    if not d: return "—"
    try:
        from datetime import datetime
        return datetime.strptime(d[:10], "%Y-%m-%d").strftime("%b %d, %Y")
    except Exception:
        return d[:10] if d else "—"


def _extract_gs_lang(goods_list: list, lang: str) -> Dict[str, str]:
    for item in goods_list:
        ga = item.get("goodAndServices", {})
        if (ga.get("language") or "").upper() == lang.upper():
            result = {}
            for entry in ga.get("goodAndServiceList", []):
                nc    = str(entry.get("niceClass", ""))
                terms = entry.get("goodsAndServices", [])
                text  = "; ".join(t.get("term", "") for t in terms if t.get("term"))
                if nc and text:
                    result[nc] = text
            return result
    return {}


def _set_borders(table, color_hex="D0D7E3"):
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top","left","bottom","right","insideH","insideV"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"),   "single")
                b.set(qn("w:sz"),    "4")
                b.set(qn("w:space"), "0")
                b.set(qn("w:color"), color_hex)
                tcBorders.append(b)
            tcPr.append(tcBorders)


def _lock_table_layout(table):
    """Force fixed table layout so Word does not auto-expand past page margins."""
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def _set_table_col_widths_cm(table, widths_cm: List[float]):
    """Set fixed widths for all columns/cells in centimeters."""
    _lock_table_layout(table)
    for i, w in enumerate(widths_cm):
        w_cm = Cm(float(w))
        table.columns[i].width = w_cm
        for row in table.rows:
            row.cells[i].width = w_cm


def _set_row_bg(row, color_hex: str):
    """Set background color for all cells in a row."""
    for cell in row.cells:
        _set_cell_bg(cell, color_hex)


def _add_section_title(doc, text: str, size: int = 13, color: RGBColor = None):
    """Add professional section title with spacing."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color or RGBColor(0x0F, 0x34, 0x60)
    r.font.name = "Arial"
    p.paragraph_format.space_after = Pt(8)
    return p


def _p(cell, text, bold=False, size=8, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, first=False):
    """Add paragraph to cell (first=True uses existing first paragraph)."""
    if first and cell.paragraphs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    if not text:
        return p
    run = p.add_run(str(text))
    run.bold       = bold
    run.font.size  = Pt(size)
    run.font.name  = "Arial"
    if color:
        run.font.color.rgb = color
    return p


STRATEGY_TEXT = {
    "title1": "Strategie de căutare - Mărci comerciale",
    "subtitle1": "Strategie de căutare aleasă:",
    "desc_title": "Descrierea acurateții",
    "desc_body": (
        "Algoritmul de căutare este construit cu o metodă statistică, folosind peste 1 milion "
        "de cazuri oficiale în care două mărci au fost găsite confuzionante de un oficial guvernamental "
        "în SUA și UE.\n\n"
        "Aceasta înseamnă că rapoartele de căutare și monitorizare acoperă toate strategiile de căutare "
        "obișnuite aplicate în căutarea manuală, cum ar fi:"
    ),
    "bullets": [
        "Identitate exactă",
        "Similaritate fonetică",
        "Similaritate ortografică și greșeli de scriere",
        "Variații de prefix, infix și sufix",
        "Similaritate între vocale și consoane",
        "Plurale și variații de rădăcină",
        "Abrevieri și acronime",
        "Alte similarități",
    ],
    "noise": (
        "Algoritmul aplică, de asemenea, tehnici unice de reducere a \"zgomotului\". Aceasta rezultă în "
        "un număr mai mic de rezultate fără a afecta calitatea generală (atingând peste 99% din "
        "potențialele conflicte)."
    ),
    "title2": "Analiza statistică a riscului - Mărci comerciale",
    "risk_intro": (
        "Clasamentul rezultatelor se bazează pe o analiză statistică a peste 1 milion de cazuri oficiale "
        "de mărci comerciale confuzionante în UE sau SUA.\n\n"
        "Cele patru \"Niveluri de risc\" indică statistic unde veți găsi potențiale conflicte."
    ),
    "levels": [
        ("Nivel 1 - Risc foarte ridicat (85-100%)",
         "Din toate conflictele, 20% au aceste tipuri de similaritati.",
         "very_high"),
        ("Nivel 2 - Risc ridicat (70-84%)",
         "Din toate conflictele in Europa sau SUA, 40% au aceste tipuri de similaritati.",
         "high"),
        ("Nivel 3 - Risc mediu (55-69%)",
         "Din toate conflictele in Europa sau SUA, 25% au aceste tipuri de similaritati.",
         "medium"),
        ("Nivel 4 - Risc scazut (40-54%)",
         "Din toate conflictele in Europa sau SUA, 15% au aceste tipuri de similaritati.",
         "low"),
    ],
}


def _word_strategy_page(doc: Document):
    """Add search strategy page to Word document."""
    BLUE = RGBColor(0x0F, 0x34, 0x60)
    GRAY = RGBColor(0x44, 0x44, 0x44)

    def h(text, size=12, bold=True, color=None):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold; r.font.size = Pt(size); r.font.name = "Arial"
        if color: r.font.color.rgb = color
        return p

    def body(text, size=9):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(size); r.font.name = "Arial"
        r.font.color.rgb = GRAY
        return p

    s = STRATEGY_TEXT
    h(s["title1"], size=13, color=BLUE)
    h(s["subtitle1"], size=10, color=BLUE)
    doc.add_paragraph()

    h(s["desc_title"], size=10, color=BLUE)
    body(s["desc_body"])
    doc.add_paragraph()

    for bullet in s["bullets"]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(bullet)
        r.font.size = Pt(9); r.font.name = "Arial"; r.font.color.rgb = GRAY

    doc.add_paragraph()
    body(s["noise"])
    doc.add_paragraph()

    h(s["title2"], size=13, color=BLUE)
    body(s["risk_intro"])
    doc.add_paragraph()

    level_colors = {
        "very_high": (RGBColor(0xC0,0x39,0x2B), "FDECEA"),
        "high":      (RGBColor(0xAF,0x4D,0x00), "FEEBCF"),
        "medium":    (RGBColor(0x9A,0x76,0x00), "FFF9DB"),
        "low":       (RGBColor(0x1E,0x84,0x49), "EAFAF1"),
    }

    for lbl, desc, lvl in s["levels"]:
        fg, bg = level_colors[lvl]
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        cell = tbl.cell(0, 0)
        _set_cell_bg(cell, bg)
        ph = cell.paragraphs[0]
        rh = ph.add_run(lbl)
        rh.bold = True; rh.font.size = Pt(9); rh.font.name = "Arial"
        rh.font.color.rgb = fg
        pd = cell.add_paragraph(desc)
        rd = pd.runs[0] if pd.runs else pd.add_run(desc)
        rd.font.size = Pt(8); rd.font.name = "Arial"; rd.font.color.rgb = GRAY
        doc.add_paragraph()


def _set_left_accent(cell, color_hex="0F3460"):
    """Adauga un chenar stanga albastru gros (ca .goods-block din web UI)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    b = OxmlElement("w:left")
    b.set(qn("w:val"),   "single")
    b.set(qn("w:sz"),    "18")
    b.set(qn("w:space"), "0")
    b.set(qn("w:color"), color_hex)
    tcBorders.append(b)
    tcPr.append(tcBorders)


def _word_trademark_card(doc, tm, page_w_cm: float = 27.1):
    BLUE   = RGBColor(0x0F, 0x34, 0x60)
    GRAY   = RGBColor(0x44, 0x44, 0x44)
    LGRAY  = RGBColor(0x88, 0x88, 0x88)
    PURPLE = RGBColor(0x8E, 0x44, 0xAD)
    RED    = RGBColor(0xC0, 0x39, 0x2B)

    sim   = tm.get("similarity") or {}
    score = sim.get("combined_score") or 0
    lvl   = _risk_level(score)
    fg    = RGBColor(*_RISK_RGB[lvl])
    bg    = _bg_hex(score)

    RISK_LABELS_W = {
        "very_high": "RISC FOARTE RIDICAT",
        "high":      "RISC RIDICAT",
        "medium":    "RISC MEDIU",
        "low":       "RISC SCĂZUT",
    }
    risk_label = RISK_LABELS_W[lvl]

    office    = tm.get("office") or ""
    office_nm = tm.get("officeName") or ""
    status    = tm.get("status") or "—"
    sl        = status.lower()
    if "registered" in sl:
        stat_txt, stat_col = "✔ Înregistrată", RGBColor(0x1E, 0x84, 0x49)
    elif "filed" in sl or "pending" in sl:
        stat_txt, stat_col = "⏳ Depusă", RGBColor(0xB7, 0x95, 0x0B)
    elif any(w in sl for w in ("expir","lapsed","cancelled","refused","withdrawn")):
        stat_txt, stat_col = "✖ Expirată/Anulată", RED
    else:
        stat_txt, stat_col = status, GRAY

    owner     = ", ".join(a.get("name","") for a in (tm.get("applicants") or []) if a.get("name")) or "—"
    app_addr  = "; ".join(a.get("address","") for a in (tm.get("applicants") or []) if a.get("address"))
    reps_w    = ", ".join((r.get("fullName") or r.get("organizationName",""))
                          for r in (tm.get("representatives") or [])
                          if r.get("fullName") or r.get("organizationName"))
    an        = tm.get("applicationNumber") or "—"
    rn        = tm.get("registrationNumber") or "—"
    app_date  = _fmt_date(tm.get("applicationDate") or "")
    reg_date  = _fmt_date(tm.get("registrationDate") or "")
    exp_date  = _fmt_date(tm.get("expiryDate") or "")
    exp_note  = " (*)" if tm.get("expiryDate") and not tm.get("expiryIsReal") else ""
    pub_date  = _fmt_date(tm.get("publicationDate") or "")
    opp_start = _fmt_date(tm.get("oppositionStartDate") or "")
    opp_end   = _fmt_date(tm.get("oppositionEndDate") or "")
    mark_feat = " · ".join(filter(None,[tm.get("markFeature") or "", tm.get("kindMark") or ""]))
    vienna    = ", ".join(tm.get("viennaCodes") or [])
    designated = ", ".join(tm.get("designatedCountries") or [])
    found_by  = tm.get("_found_by") or ""

    # NICE sortate crescator
    nice_detailed_w = sorted(
        tm.get("niceDetailed") or [],
        key=lambda nd: int(str(nd.get("class","0"))) if str(nd.get("class","0")).isdigit() else 0
    )
    if nice_detailed_w:
        classes_str = "  |  ".join(
            f"Cls {nd['class']} — {nd.get('short','')}" for nd in nice_detailed_w)
    else:
        classes_str = "  ".join(
            f"Cls {c}" for c in sorted(tm.get("niceClass") or [],
                                       key=lambda x: int(x) if str(x).isdigit() else 0))

    # ── Dimensiuni coloane dinamice (cu marja interna de siguranta) ─────
    CARD_W  = max(page_w_cm - 0.8, 21.8)
    STRIP_W = 0.35
    LOGO_W  = 2.20
    SCORE_W = 4.10
    INFO_W  = CARD_W - STRIP_W - LOGO_W - SCORE_W

    # ── Card: [strip | logo | info | score] ─────────────────────────────
    card = doc.add_table(rows=1, cols=4)
    card.style = "Table Grid"
    _set_table_col_widths_cm(card, [STRIP_W, LOGO_W, INFO_W, SCORE_W])
    sc, lc, ic, rc = card.cell(0,0), card.cell(0,1), card.cell(0,2), card.cell(0,3)

    # Strip (culoare risc)
    _set_cell_bg(sc, bg)
    sc.paragraphs[0].add_run("")

    # Logo
    img_bytes = _fetch_image_bytes(tm.get("imageUrl"), size=(65, 65))
    p_logo = lc.paragraphs[0]; p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if img_bytes:
        try:
            p_logo.add_run().add_picture(io.BytesIO(img_bytes), width=Cm(1.9))
        except Exception:
            r = p_logo.add_run("TM"); r.font.size = Pt(20); r.font.color.rgb = RGBColor(0xBB,0xBB,0xBB)
    else:
        r = p_logo.add_run("TM"); r.font.size = Pt(20); r.font.name = "Arial"
        r.font.color.rgb = RGBColor(0xBB,0xBB,0xBB)

    # Info column
    _p(ic, tm.get("tmName") or "—", bold=True, size=13, color=RGBColor(0x1a,0x1a,0x2e), first=True)

    p_meta = ic.add_paragraph()
    r_off = p_meta.add_run(f" {office} ")
    r_off.bold = True; r_off.font.size = Pt(8.5); r_off.font.name = "Arial"; r_off.font.color.rgb = BLUE
    if office_nm:
        r_onm = p_meta.add_run(f"  {office_nm}  ")
        r_onm.font.size = Pt(7.5); r_onm.font.name = "Arial"; r_onm.font.color.rgb = LGRAY
    r_st = p_meta.add_run(stat_txt)
    r_st.font.size = Pt(8.5); r_st.font.name = "Arial"; r_st.font.color.rgb = stat_col

    p_own = ic.add_paragraph()
    rl = p_own.add_run("Titular: "); rl.bold = True; rl.font.size = Pt(9); rl.font.name = "Arial"
    rv = p_own.add_run(owner);       rv.font.size = Pt(9); rv.font.name = "Arial"

    _p(ic, f"Nr. marcă: {rn}   |   Nr. depozit: {an}", size=8, color=GRAY)

    p_dates = ic.add_paragraph()
    for lbl_d, val_d, col_d in [
        ("Depus: ",  app_date or "—", GRAY),
        ("   Înreg.: ", reg_date or "—", GRAY),
        ("   Expiră: ", f"{exp_date}{exp_note}" if exp_date else "—", RED if exp_date else GRAY),
    ]:
        rl2 = p_dates.add_run(lbl_d); rl2.font.size = Pt(7.5); rl2.font.name = "Arial"; rl2.font.color.rgb = LGRAY
        rv2 = p_dates.add_run(val_d); rv2.font.size = Pt(7.5); rv2.font.name = "Arial"; rv2.font.color.rgb = col_d
        if col_d == RED: rv2.bold = True

    _p(ic, classes_str, size=8, color=BLUE)

    # Score column
    _set_cell_bg(rc, bg)
    p_sc = rc.paragraphs[0]; p_sc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sc = p_sc.add_run(f"{score}%\n")
    r_sc.bold = True; r_sc.font.size = Pt(28); r_sc.font.name = "Arial"; r_sc.font.color.rgb = fg
    r_rl = p_sc.add_run(risk_label)
    r_rl.bold = True; r_rl.font.size = Pt(7.5); r_rl.font.name = "Arial"; r_rl.font.color.rgb = fg

    p_sc2 = rc.add_paragraph(); p_sc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_score = sim.get("textual_score") or 0
    p_score = sim.get("phonetic_score") or 0
    for lbl_s, col_s in [
        (f"📝 {t_score}%  textual", GRAY),
        (f"🔊 {p_score}%  fonetic", PURPLE),
        (f"Jaro-W: {sim.get('jaro_winkler',0)}%", LGRAY),
        (f"Lev.: {sim.get('levenshtein_distance',0)} car.", LGRAY),
    ]:
        r = p_sc2.add_run(lbl_s + "\n"); r.font.size = Pt(7); r.font.name = "Arial"; r.font.color.rgb = col_s

    _set_borders(card)

    # ── Detalii suplimentare (2 coloane egale) ───────────────────────────
    extra_w = [
        ("Data publicare (450)", pub_date),
        ("Perioada opozitie",    f"{opp_start} – {opp_end}" if opp_start else ""),
        ("Natura marcii (550)",  mark_feat),
        ("Coduri Vienna (531)",  vienna),
        ("Tari desemnate",       designated),
        ("Reprezentant (740)",   reps_w),
        ("Adresa titular",       app_addr),
        ("ST13",                 tm.get("ST13") or ""),
        ("Gasit prin varianta",  found_by),
    ]
    active_extra = [(l, v) for l, v in extra_w if v]
    if active_extra:
        det_col = CARD_W / 3
        det_tbl = doc.add_table(rows=1, cols=3); det_tbl.style = "Table Grid"
        _set_table_col_widths_cm(det_tbl, [det_col, det_col, det_col])
        for ci in range(3):
            _set_cell_bg(det_tbl.cell(0, ci), "FAFBFD")
        targets = [det_tbl.cell(0, 0), det_tbl.cell(0, 1), det_tbl.cell(0, 2)]
        _p(targets[0], "Detalii suplimentare", bold=True, size=8, color=BLUE, first=True)
        _p(targets[1], "", first=True)
        _p(targets[2], "", first=True)
        for idx, (lbl_d, val_d) in enumerate(active_extra):
            target = targets[idx % 3]
            pd = target.add_paragraph()
            rl3 = pd.add_run(f"{lbl_d}: "); rl3.bold = True; rl3.font.size = Pt(7); rl3.font.name = "Arial"; rl3.font.color.rgb = LGRAY
            rv3 = pd.add_run(str(val_d));   rv3.font.size = Pt(7.5); rv3.font.name = "Arial"
        _set_borders(det_tbl)

    # ── G&S: blocuri cu chenar sortate crescator (fara repetare short) ───
    all_cls_w: dict = {}
    for g in (tm.get("goodAndServices") or []):
        nc = str(g.get("niceClass") or "")
        if nc and nc.isdigit():
            if nc not in all_cls_w: all_cls_w[nc] = {"text":"","short":"","desc":""}
            all_cls_w[nc]["text"]  = g.get("goodsAndServices") or ""
            all_cls_w[nc]["short"] = g.get("niceShort") or ""
    for nd in nice_detailed_w:
        nc = str(nd.get("class",""))
        if nc and nc.isdigit():
            if nc not in all_cls_w: all_cls_w[nc] = {"text":"","short":"","desc":""}
            if not all_cls_w[nc]["short"]: all_cls_w[nc]["short"] = nd.get("short") or ""
            all_cls_w[nc]["desc"] = nd.get("description") or ""

    if all_cls_w:
        p_gs_h = doc.add_paragraph()
        r_gs_h = p_gs_h.add_run("CLASIFICARE INTERNATIONALA NISA / PRODUSE SI SERVICII")
        r_gs_h.bold = True; r_gs_h.font.size = Pt(9); r_gs_h.font.name = "Arial"; r_gs_h.font.color.rgb = BLUE
        p_gs_h.paragraph_format.space_before = Pt(4); p_gs_h.paragraph_format.space_after = Pt(3)

        for nc in sorted(all_cls_w.keys(), key=lambda x: int(x)):
            info = all_cls_w[nc]
            gs_t = doc.add_table(rows=1, cols=1); gs_t.style = "Table Grid"
            _set_table_col_widths_cm(gs_t, [CARD_W])
            gs_c2 = gs_t.cell(0,0)
            _set_cell_bg(gs_c2, "FFFFFF"); _set_left_accent(gs_c2, "0F3460")

            _p(gs_c2, f"Clasa {nc}", bold=True, size=8.5, color=BLUE, first=True)
            if info["short"]:
                _p(gs_c2, info["short"], size=8, color=GRAY)
            if info["text"]:
                _p(gs_c2, info["text"], size=8, color=RGBColor(0x33,0x33,0x33))
            if info["desc"]:
                pd2 = gs_c2.add_paragraph()
                rd2 = pd2.add_run(info["desc"])
                rd2.font.size = Pt(7); rd2.font.name = "Arial"; rd2.italic = True
                rd2.font.color.rgb = RGBColor(0xAA,0xAA,0xAA)
            _set_borders(gs_t)

            sp2 = doc.add_paragraph()
            sp2.paragraph_format.space_before = Pt(0); sp2.paragraph_format.space_after = Pt(3)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0); sp.paragraph_format.space_after = Pt(8)


def build_word(query: str, nice_classes: List[str], offices: List[str],
               results: List[Dict], similar: List[Dict] = None,
               expired_conflicts: List[Dict] = None, expired_similar: List[Dict] = None) -> bytes:
    from datetime import datetime as dt
    from docx.enum.section import WD_ORIENT

    # A4 landscape: 29.7 × 21 cm, margini 1.3 cm.
    # We keep a small safety margin inside the usable width so Word tables
    # do not spill outside the printable area.
    PAGE_W_CM = 26.2
    CONTENT_W_CM = PAGE_W_CM - 0.6
    MARGIN    = Cm(1.3)

    doc = Document()
    for sec in doc.sections:
        sec.orientation   = WD_ORIENT.LANDSCAPE
        sec.page_width    = Cm(29.7)
        sec.page_height   = Cm(21.0)
        sec.top_margin    = MARGIN
        sec.bottom_margin = MARGIN
        sec.left_margin   = MARGIN
        sec.right_margin  = MARGIN

    _add_export_brand_header_word(doc, query)

    active_results = _sort_export_results((results or []) + (similar or []), offices)
    expired_results = _sort_export_results((expired_conflicts or []) + (expired_similar or []), offices)
    all_results = active_results + expired_results

    risk_buckets = _risk_buckets(all_results)
    very_high = risk_buckets["very_high"]
    high      = risk_buckets["high"]
    medium    = risk_buckets["medium"]
    low       = risk_buckets["low"]

    risky_count  = len(very_high) + len(high)
    similar_count = len(medium) + len(low)
    safe = risky_count == 0

    BLUE  = RGBColor(0x0F, 0x34, 0x60)
    GRAY  = RGBColor(0x44, 0x44, 0x44)

    # Distributie pe oficii
    geo_counts = {}
    for tm in all_results:
        o = tm.get("office") or tm.get("tmOffice") or "?"
        geo_counts[o] = geo_counts.get(o, 0) + 1
    geo_sorted = sorted(geo_counts.items(), key=lambda x: x[1], reverse=True)

    # ─── COVER PAGE (dashboard ca in UI) ─────────────────────────────
    # Titlu aplicatie
    p_app = doc.add_paragraph()
    r_app = p_app.add_run("Verificare Disponibilitate Marcă")
    r_app.font.size = Pt(10); r_app.font.name = "Arial"; r_app.font.color.rgb = GRAY

    # Numele marcii cautat
    p_q = doc.add_paragraph()
    r_q = p_q.add_run(query)
    r_q.bold = True; r_q.font.size = Pt(22); r_q.font.name = "Arial"
    r_q.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    p_q.paragraph_format.space_after = Pt(8)

    cover_divider = doc.add_table(rows=1, cols=1)
    cover_divider.style = "Table Grid"
    _set_table_col_widths_cm(cover_divider, [CONTENT_W_CM])
    _set_cell_bg(cover_divider.cell(0, 0), "0F3460")
    cover_divider.cell(0, 0).paragraphs[0].add_run("")
    _set_borders(cover_divider, "0F3460")
    doc.add_paragraph()

    # Badges — tabel 3 coloane x 2 randuri
    def _wbadge(cell, label, value, bg_hex, fg_rgb, bold_val=True):
        _set_cell_bg(cell, bg_hex)
        p_lbl = cell.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.size = Pt(7.5); r_lbl.font.name = "Arial"
        r_lbl.font.color.rgb = fg_rgb
        p_val = cell.add_paragraph()
        r_val = p_val.add_run(str(value))
        r_val.bold = bold_val; r_val.font.size = Pt(14 if bold_val else 9)
        r_val.font.name = "Arial"; r_val.font.color.rgb = fg_rgb
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

    badge_tbl = doc.add_table(rows=2, cols=3)
    badge_tbl.style = "Table Grid"
    badge_col = CONTENT_W_CM / 3
    _set_table_col_widths_cm(badge_tbl, [badge_col, badge_col, badge_col])

    RED   = RGBColor(0xC0, 0x39, 0x2B)
    GREEN = RGBColor(0x1E, 0x84, 0x49)
    ORG   = RGBColor(0x85, 0x64, 0x04)

    _wbadge(badge_tbl.cell(0,0), "Total găsite",                 len(all_results),  "E8F0FB", BLUE)
    _wbadge(badge_tbl.cell(0,1), "Risc ridicat / f.ridicat >=70%", risky_count,
            "FDECEA" if not safe else "EAFAF1", RED if not safe else GREEN)
    _wbadge(badge_tbl.cell(0,2), "Risc mediu / scăzut 40-70%", similar_count,     "FFF3CD", ORG)
    _wbadge(badge_tbl.cell(1,0), "Clase NICE",  ", ".join(nice_classes),            "E8F0FB", BLUE, bold_val=False)
    _wbadge(badge_tbl.cell(1,1), "Mărci expirate",              str(len(expired_results)), "FDECEA", RED)
    _wbadge(badge_tbl.cell(1,2), "Data raport",
            date.today().strftime("%d.%m.%Y"),                                      "F2F3F4", GRAY, bold_val=False)
    doc.add_paragraph()

    # Distribuție pe oficii
    if geo_sorted:
        p_geo_title = doc.add_paragraph()
        r_geo = p_geo_title.add_run("Distribuție pe oficii")
        r_geo.bold = True; r_geo.font.size = Pt(9); r_geo.font.name = "Arial"
        r_geo.font.color.rgb = BLUE
        p_geo_title.paragraph_format.space_after = Pt(4)

        OFFICE_NAMES_W = {
            "EM":"EUIPO (UE)", "WO":"WIPO (Intl)", "RO":"OSIM Romania",
            "DE":"DPMA Germania","FR":"INPI Franta","IT":"UIBM Italia",
            "ES":"OEPM Spania","PL":"UPRP Polonia","BG":"BPO Bulgaria",
            "HU":"HIPO Ungaria","CZ":"IPO Cehia","AT":"APO Austria",
            "NL":"BOIP Olanda","BE":"BOIP Belgia","PT":"INPI Portugalia",
            "SE":"PRV Suedia","DK":"DKPTO Danemarca","GB":"UKIPO Marea Britanie",
        }
        geo_max = geo_sorted[0][1] if geo_sorted else 1
        geo_tbl_w = doc.add_table(rows=1, cols=3)
        geo_tbl_w.style = "Table Grid"
        geo_w = [2.0, CONTENT_W_CM - 5.0, 3.0]
        _set_table_col_widths_cm(geo_tbl_w, geo_w)
        for ci, hdr in enumerate(["Cod", "Oficiu", "Mărci"]):
            c = geo_tbl_w.cell(0, ci)
            _set_cell_bg(c, "0F3460")
            r = c.paragraphs[0].add_run(hdr)
            r.bold = True; r.font.size = Pt(8); r.font.name = "Arial"
            r.font.color.rgb = RGBColor(255, 255, 255)

        for ri, (code, cnt) in enumerate(geo_sorted, 1):
            row_w = geo_tbl_w.add_row()
            row_w.height = Cm(0.55)
            is_max = cnt == geo_max
            fg_w = RED if is_max else BLUE
            _set_cell_bg(row_w.cells[0], "FDECEA" if is_max else "F7F9FC")
            r0 = row_w.cells[0].paragraphs[0].add_run(code)
            r0.bold = True; r0.font.size = Pt(9); r0.font.name = "Arial"; r0.font.color.rgb = fg_w
            r1 = row_w.cells[1].paragraphs[0].add_run(OFFICE_NAMES_W.get(code, code))
            r1.font.size = Pt(8); r1.font.name = "Arial"; r1.font.color.rgb = GRAY
            r2 = row_w.cells[2].paragraphs[0].add_run(str(cnt))
            r2.bold = True; r2.font.size = Pt(10); r2.font.name = "Arial"; r2.font.color.rgb = fg_w
            row_w.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_borders(geo_tbl_w)
        doc.add_paragraph()

    doc.add_page_break()

    # ─── RESULTS PAGES ─────────────────────────────────────────────────
    _add_section_title(doc, "Rezultate analiza similaritate")
    _add_section_title(doc, f"Mărci active ({len(active_results)})", size=16)

    sub_p = doc.add_paragraph()
    sub_p.add_run(query)
    sub_p.runs[0].font.size = Pt(10)
    sub_p.runs[0].font.color.rgb = RGBColor(0x44,0x44,0x44)
    sub_p.runs[0].font.name = "Arial"
    sub_p.paragraph_format.space_after = Pt(8)

    if not active_results and not expired_results:
        p_empty = doc.add_paragraph("Niciun conflict detectat.")
        p_empty.runs[0].font.color.rgb = RGBColor(0x1E,0x84,0x49)
    else:
        for tm in active_results:
            _word_trademark_card(doc, tm, page_w_cm=PAGE_W_CM)

    # ─── SUMMARY PAGE ──────────────────────────────────────────────────
    doc.add_page_break()
    _add_section_title(doc, "Sumar rezultate")

    risk_groups = [
        (very_high, "Risc foarte ridicat  (>= 90%)", "C0392B", "FDECEA"),
        (high,      "Risc ridicat  (75-89%)",         "AF4D00", "FEEBCF"),
        (medium,    "Risc mediu  (60-74%)",            "9A7600", "FFF9DB"),
        (low,       "Risc scăzut  (45-59%)",           "1E8449", "EAFAF1"),
    ]

    for grp, lbl, fg_hex, bg_hex2 in risk_groups:
        fgc = RGBColor(int(fg_hex[0:2],16), int(fg_hex[2:4],16), int(fg_hex[4:6],16))
        sh = doc.add_table(rows=1, cols=1); sh.style = "Table Grid"
        _set_table_col_widths_cm(sh, [CONTENT_W_CM])
        shc = sh.cell(0,0); _set_cell_bg(shc, bg_hex2)
        shr = shc.paragraphs[0].add_run(f"  {lbl}  -  {len(grp)} mărci")
        shr.bold = True; shr.font.size = Pt(10); shr.font.name = "Arial"; shr.font.color.rgb = fgc
        doc.add_paragraph()

        if not grp:
            doc.add_paragraph("  Nicio marcă.").runs[0].font.size = Pt(9)
            doc.add_paragraph(); continue

        st2 = doc.add_table(rows=1, cols=6); st2.style = "Table Grid"
        sh2 = ["#", "Denumire marcă", "Oficiu", "Titular", "Status", "Scor"]
        summary_scale = CONTENT_W_CM / 27.1
        sw2_cm = [v * summary_scale for v in (0.8, 7.5, 3.3, 8.5, 3.8, 3.2)]
        _set_table_col_widths_cm(st2, sw2_cm)
        for ci,ch in enumerate(sh2):
            c = st2.cell(0,ci); _set_cell_bg(c,"0F3460")
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(ch); r.bold = True; r.font.size = Pt(8); r.font.name = "Arial"
            r.font.color.rgb = RGBColor(255,255,255)

        for ri,tm in enumerate(grp,1):
            score2    = tm.get("similarity",{}).get("combined_score",0)
            applicant = ", ".join(a.get("name","") for a in tm.get("applicants",[]) if a.get("name")) or "—"
            row2      = st2.add_row()
            row2.height = Cm(0.6)
            vals      = [str(ri), tm.get("tmName","—"), tm.get("office","—"),
                         applicant, tm.get("status") or "—", f"{score2}%"]
            for ci,val in enumerate(vals):
                c = row2.cells[ci]
                if ri%2==0: _set_cell_bg(c,"F7F9FC")
                p = c.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci in (0,2,4,5) else WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(val); r.font.size = Pt(8); r.font.name = "Arial"
                if ci==5:
                    r.bold=True; r.font.color.rgb=RGBColor(*_RISK_RGB[_risk_level(score2)])
        _set_borders(st2)
        doc.add_paragraph()

    # ─── CONCLUSIONS PAGE ──────────────────────────────────────────────
    doc.add_page_break()
    _add_section_title(doc, "Concluzii și recomandări")

    high_risk_count = len(very_high) + len(high)
    medium_risk_count = len(medium)
    minimal_risk_count = len(low)

    if high_risk_count:
        risk_heading = "Concluzie privind riscul relativ: risc semnificativ"
        risk_color = RGBColor(0xC0, 0x39, 0x2B)
        risk_body = (
            f"Au fost identificate {high_risk_count} mărci cu risc ridicat sau foarte ridicat, "
            f"{medium_risk_count} mărci cu risc mediu și {minimal_risk_count} mărci cu risc minim/scăzut. "
            "În raport cu motivele relative de refuz, acest rezultat indică o probabilitate relevantă de "
            "opoziție sau de refuz întemeiat pe existența unor drepturi anterioare, în special atunci când "
            "se cumulează similitudinea semnului cu proximitatea produselor sau serviciilor revendicate."
        )
    elif medium_risk_count:
        risk_heading = "Concluzie privind riscul relativ: risc moderat"
        risk_color = RGBColor(0xAF, 0x4D, 0x00)
        risk_body = (
            f"Au fost identificate {high_risk_count} mărci cu risc ridicat sau foarte ridicat, "
            f"{medium_risk_count} mărci cu risc mediu și {minimal_risk_count} mărci cu risc minim/scăzut. "
            "În raport cu motivele relative de refuz, rezultatele indică necesitatea unei evaluări juridice "
            "suplimentare a mărcilor aflate în zona medie, deoarece riscul de confuzie depinde de impresia de "
            "ansamblu a semnelor și de gradul de apropiere dintre produsele sau serviciile vizate."
        )
    else:
        risk_heading = "Concluzie privind riscul relativ: risc redus"
        risk_color = RGBColor(0x1E, 0x84, 0x49)
        risk_body = (
            f"Au fost identificate {high_risk_count} mărci cu risc ridicat sau foarte ridicat, "
            f"{medium_risk_count} mărci cu risc mediu și {minimal_risk_count} mărci cu risc minim/scăzut. "
            "Pe baza conflictelor identificate, motivele relative de refuz par limitate; totuși, disponibilitatea "
            "juridică finală trebuie confirmată prin analiza individuală a drepturilor anterioare relevante."
        )

    conclusion_blocks = [
        (risk_heading, risk_body, risk_color),
        (
            "Opinia privind motivele absolute de refuz",
            "Prezentul raport are în principal o funcție de cercetare a conflictelor relative și nu poate stabili "
            "în mod definitiv incidența motivelor absolute de refuz. În practica oficiilor și a jurisprudenței "
            "europene, examinarea motivelor absolute vizează în special caracterul distinctiv, eventuala "
            "descriptivitate, caracterul uzual, posibilul caracter înșelător și conformitatea semnului cu ordinea "
            "publică. În consecință, marca trebuie analizată separat și prin raportare exactă la lista produselor "
            "și serviciilor revendicate.",
            BLUE,
        ),
        (
            "Opinia privind validitatea mărcii",
            "Condiția validității unei viitoare înregistrări depinde cumulativ de absența impedimentelor absolute, "
            "de inexistența unor drepturi anterioare opozabile și de formularea adecvată a specificației de produse "
            "și servicii. Chiar și în ipoteza unui risc relativ redus, validitatea nu poate fi considerată automat "
            "îndeplinită fără verificarea completă a cadrului juridic aplicabil pe teritoriile selectate.",
            BLUE,
        ),
        (
            "Opinia privind distinctivitatea",
            "Distinctivitatea trebuie apreciată din perspectiva publicului relevant și în raport direct cu produsele "
            "sau serviciile pentru care se solicită protecția. În mod obișnuit, semnele fanteziste sau arbitrare au "
            "o forță distinctivă mai ridicată, în timp ce semnele descriptive, laudative ori slab individualizante "
            "sunt mai expuse obiecțiilor la examinare și beneficiază de o protecție mai restrânsă. Dacă denumirea "
            "propusă evocă în mod imediat natura, calitatea, destinația sau alte caracteristici ale produselor sau "
            "serviciilor, se recomandă o reevaluare a semnului înainte de depunere.",
            BLUE,
        ),
        (
            "Recomandare finală",
            "Rezultatele acestui raport trebuie utilizate ca instrument de triere a riscului, nu ca opinie juridică "
            "definitivă. Pentru depunere, este recomandată o analiză specializată privind opozabilitatea drepturilor "
            "anterioare, formularea claselor NICE și sustenabilitatea mărcii sub aspectul distinctivității și al "
            "motivelor absolute de refuz.",
            BLUE,
        ),
    ]

    for title, body, color in conclusion_blocks:
        p_title = doc.add_paragraph()
        r_title = p_title.add_run(title)
        r_title.bold = True
        r_title.font.size = Pt(10.5)
        r_title.font.name = "Arial"
        r_title.font.color.rgb = color
        p_title.paragraph_format.space_after = Pt(4)

        p_body = doc.add_paragraph(body)
        for run in p_body.runs:
            run.font.size = Pt(9)
            run.font.name = "Arial"
            run.font.color.rgb = GRAY
        p_body.paragraph_format.space_after = Pt(8)

    _add_section_title(doc, f"Mărci expirate ({len(expired_results)})", size=16, color=RED)
    p_exp = doc.add_paragraph()
    r_exp = p_exp.add_run("Aceste mărci sunt listate separat deoarece au statut expirat/inactiv.")
    r_exp.font.size = Pt(8)
    r_exp.font.name = "Arial"
    r_exp.font.color.rgb = GRAY
    p_exp.paragraph_format.space_after = Pt(6)

    sh_exp = doc.add_table(rows=1, cols=1); sh_exp.style = "Table Grid"
    _set_table_col_widths_cm(sh_exp, [CONTENT_W_CM])
    shc_exp = sh_exp.cell(0,0); _set_cell_bg(shc_exp, "FDECEA")
    shr_exp = shc_exp.paragraphs[0].add_run(f"  Mărci expirate  -  {len(expired_results)} mărci")
    shr_exp.bold = True; shr_exp.font.size = Pt(10); shr_exp.font.name = "Arial"; shr_exp.font.color.rgb = RED
    doc.add_paragraph()

    if expired_results:
        st_exp = doc.add_table(rows=1, cols=5); st_exp.style = "Table Grid"
        exp_headers = ["#", "Denumire marcă", "Oficiu", "Status", "Scor"]
        exp_widths = [0.8, 9.2, 3.2, 9.0, 2.4]
        _set_table_col_widths_cm(st_exp, exp_widths)
        for ci, ch in enumerate(exp_headers):
            c = st_exp.cell(0, ci); _set_cell_bg(c, "C0392B")
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(ch); r.bold = True; r.font.size = Pt(8); r.font.name = "Arial"; r.font.color.rgb = RGBColor(255,255,255)
        for ri, tm in enumerate(expired_results, 1):
            score2 = tm.get("similarity",{}).get("combined_score",0)
            row = st_exp.add_row()
            row.height = Cm(0.6)
            vals = [str(ri), tm.get("tmName","—"), tm.get("office","—"), tm.get("status") or "—", f"{score2}%"]
            for ci, val in enumerate(vals):
                c = row.cells[ci]
                if ri % 2 == 0: _set_cell_bg(c, "FFF5F5")
                p = c.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci in (0,2,4) else WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(val); r.font.size = Pt(8); r.font.name = "Arial"
                if ci == 4:
                    r.bold = True; r.font.color.rgb = RED
        _set_borders(st_exp)
        doc.add_paragraph()
    else:
        p_noexp = doc.add_paragraph("Nicio marcă expirată.")
        p_noexp.runs[0].font.size = Pt(9)
        doc.add_paragraph()

    # Footer note
    doc.add_paragraph()
    p_note = doc.add_paragraph()
    r_note = p_note.add_run(
        "Notă: Acest raport este generat automat și nu constituie consultanță juridică oficială. "
        "Datele provin din baze de date publice (TMview, EUIPO, WIPO) și sunt actualizate periodic."
    )
    r_note.font.size = Pt(8)
    r_note.font.italic = True
    r_note.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p_note.paragraph_format.space_after = Pt(0)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
